import argparse
import re
from dataclasses import dataclass
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


QUESTION_START_RE = re.compile(
    r"(?m)^\s*(?:"
    r"प्रश्न\s*\.?\s*(\d+)\b"
    r"|(?:Q|Question)\s*\.?\s*([0-9A-Za-z]+)\b"
    r")"
)

# Option labels can appear like:
#   (a) / (A)
#   a) / A)
#   a. / A.
#   a: / a-  (sometimes)
OPT_PAREN_RE = re.compile(r"(?<![A-Za-z0-9])\(\s*([a-dA-D])\s*\)\s*")
OPT_DELIM_RE = re.compile(r"(?<![A-Za-z0-9])([a-dA-D])\s*(?:[.)]|[:\-])\s*")

ANSWER_RE = re.compile(
    r"(?im)\b(?:उत्तर|Answer)\b\s*[:\-]?\s*"
    r"(?:\(\s*([a-dA-D])\s*\)|([a-dA-D]))\b"
)

EXPL_HEADING_RE = re.compile(
    r"(?im)\b(?:स्रोत\s*/\s*स्पष्टीकरण|स्पष्टीकरण|Explanation)\b\s*[:\-]?"
)
FACT_HEADING_RE = re.compile(
    r"(?im)\b(?:अन्य\s*महत्वपूर्ण\s*तथ्य|Other\s+Important\s+Facts|Additional\s+Facts)\b\s*[:\-]?"
)


@dataclass
class Option:
    key: str  # e.g. "(a)"
    text: str  # preserved internal newlines


@dataclass
class Question:
    no: str
    question: str
    options: List[Option]
    correct: str  # e.g. "(c)" or ""
    explanation: str
    facts: str


def normalize_letter(letter: str) -> str:
    return letter.lower()


def set_line_spacing(
    para,
    line_pts: float = 10,
    space_after_pts: float = 2,
    space_before_pts: float = 0,
):
    """
    Use Word XML to force exact line spacing.
    """
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(line_pts * 20)))
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:after"), str(int(space_after_pts * 20)))
    spacing.set(qn("w:before"), str(int(space_before_pts * 20)))
    pPr.append(spacing)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        tag = OxmlElement(f"w:{edge}")
        for key, val in attrs.items():
            tag.set(qn(f"w:{key}"), val)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_paragraph_with_text(cell, text: str, *, bold: bool, font_name: str, font_size: float, line_pts: float, after: float, before: float = 0):
    """
    Adds a paragraph and preserves internal '\n' as Word line breaks.
    """
    p = cell.add_paragraph()
    set_line_spacing(p, line_pts=line_pts, space_after_pts=after, space_before_pts=before)
    # Preserve internal newlines from extracted text
    parts = text.split("\n")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = font_name
        if i < len(parts) - 1:
            p.add_run().add_break()
    return p


def split_questions_by_paragraphs(paragraph_texts: List[str]) -> List[Tuple[str, List[str]]]:
    """
    Returns list of (question_no, block_paragraphs).
    Splits only when a paragraph starts with a question marker.
    """
    blocks: List[Tuple[str, List[str]]] = []
    current_no: Optional[str] = None
    current_paras: List[str] = []

    for raw in paragraph_texts:
        text = raw.strip()
        if not text:
            continue
        m = QUESTION_START_RE.search(text)
        if m:
            # Flush previous
            if current_no is not None:
                blocks.append((current_no, current_paras))
            current_no = m.group(1) or m.group(2) or ""
            current_paras = [raw]
        else:
            if current_no is None:
                # Skip leading garbage until the first question marker
                continue
            current_paras.append(raw)

    if current_no is not None:
        blocks.append((current_no, current_paras))
    return blocks


def find_option_spans(block_text: str) -> List[Tuple[str, int, int]]:
    """
    Returns list of (letter, label_start_index, content_start_index).
    content_start_index is where the option text begins (after the label).
    We'll compute actual ends later by looking at next content_start_index.
    """
    matches: List[Tuple[str, int, int]] = []

    for m in OPT_PAREN_RE.finditer(block_text):
        letter = normalize_letter(m.group(1))
        matches.append((letter, m.start(), m.end()))

    for m in OPT_DELIM_RE.finditer(block_text):
        letter = normalize_letter(m.group(1))
        matches.append((letter, m.start(), m.end()))

    # De-dup overlaps: if two matches overlap heavily, keep the longer one.
    matches_sorted = sorted(matches, key=lambda x: (x[1], -(x[2] - x[1])))
    chosen: List[Tuple[str, int, int]] = []
    last_end = -1
    for letter, start, end in matches_sorted:
        if start < last_end:
            # overlap: skip the overlapping one
            continue
        chosen.append((letter, start, end))
        last_end = end

    # Final structure: (letter, label_start, content_start)
    spans = [(letter, start, end) for (letter, start, end) in chosen]
    return spans


def extract_options(block_text: str) -> Tuple[List[Option], int]:
    """
    Extract options from block_text.
    Returns (options, first_option_content_start_index).
    """
    spans = find_option_spans(block_text)
    if not spans:
        return [], 0

    # Compute each option end as next option content start
    options: List[Option] = []
    for i, (letter, _label_start, content_start) in enumerate(spans):
        next_content_start = spans[i + 1][2] if i + 1 < len(spans) else len(block_text)
        option_text = block_text[content_start:next_content_start].strip()
        # Preserve internal newlines; only trim outer whitespace.
        options.append(Option(key=f"({letter})", text=option_text))

    first_label_start = spans[0][1]
    return options, first_label_start


def strip_question_marker(block_text: str, question_no: str) -> str:
    """
    Remove only the leading question marker part; keep the rest.
    """
    # Hindi: "प्रश्न <no> ..."
    t = block_text
    t = re.sub(
        rf"^\s*(?:प्रश्न\s*\.?\s*{re.escape(question_no)}\b)\s*",
        "",
        t,
        flags=re.IGNORECASE,
    )
    # English: "Q.<no> ..." or "Question <no> ..."
    t = re.sub(
        rf"^\s*(?:(?:Q|Question)\s*\.?\s*{re.escape(question_no)}\b)\s*",
        "",
        t,
        flags=re.IGNORECASE,
    )
    return t.strip()


def extract_answer(block_text: str) -> str:
    m = ANSWER_RE.search(block_text)
    if not m:
        return ""
    letter = m.group(1) or m.group(2) or ""
    if not letter:
        return ""
    return f"({normalize_letter(letter)})"


def extract_explanation_and_facts(block_text: str) -> Tuple[str, str]:
    expl_m = EXPL_HEADING_RE.search(block_text)
    fact_m = FACT_HEADING_RE.search(block_text)

    explanation = ""
    facts = ""

    if expl_m:
        expl_start = expl_m.end()
        if fact_m:
            explanation = block_text[expl_start:fact_m.start()].strip()
            facts = block_text[fact_m.end():].strip()
        else:
            explanation = block_text[expl_start:].strip()
    elif fact_m:
        facts = block_text[fact_m.end():].strip()
    else:
        # No explicit headings: best-effort.
        # Prefer starting after the answer line to avoid duplicating question/options.
        ans_m = ANSWER_RE.search(block_text)
        if ans_m:
            explanation = block_text[ans_m.end() :].strip()
        else:
            explanation = block_text.strip()

    return explanation, facts


def parse_question_block(question_no: str, block_paragraphs: List[str]) -> Question:
    block_text = "\n".join([p for p in block_paragraphs if p is not None]).strip()

    # Answer/explanation extraction first (they don't depend on options)
    correct = extract_answer(block_text)
    explanation, facts = extract_explanation_and_facts(block_text)

    # Options + question text: determine question prefix as everything before first option label.
    cleaned = strip_question_marker(block_text, question_no)
    options, first_option_start = extract_options(cleaned)

    if options:
        question_text = cleaned[:first_option_start].strip()
    else:
        # Fallback: question is everything up to answer/explanation headings if present.
        cut = len(cleaned)
        a_m = ANSWER_RE.search(cleaned)
        if a_m:
            cut = min(cut, a_m.start())
        exp_m = EXPL_HEADING_RE.search(cleaned)
        if exp_m:
            cut = min(cut, exp_m.start())
        fact_m = FACT_HEADING_RE.search(cleaned)
        if fact_m:
            cut = min(cut, fact_m.start())
        question_text = cleaned[:cut].strip()

    return Question(
        no=str(question_no),
        question=question_text,
        options=options,
        correct=correct,
        explanation=explanation,
        facts=facts,
    )


def is_option_short(opt: Option, *, inline_threshold_chars: int = 55) -> bool:
    """
    Decide if an option should be rendered inline (paired).
    We avoid pairing if the text is long or contains line breaks.
    """
    if not opt.text:
        return True
    if "\n" in opt.text:
        return False
    compact = re.sub(r"\s+", "", opt.text)
    return len(compact) <= inline_threshold_chars


def layout_option_paragraphs(options: List[Option]) -> List[List[Option]]:
    """
    Returns a list of rows.
    Each row is either [opt] or [opt1,opt2] for pairing.
    """
    rows: List[List[Option]] = []
    i = 0
    while i < len(options):
        if i + 1 < len(options):
            o1, o2 = options[i], options[i + 1]
            if is_option_short(o1) and is_option_short(o2):
                # Use a combined length heuristic too.
                combined = len(re.sub(r"\s+", "", o1.text)) + len(re.sub(r"\s+", "", o2.text))
                if combined <= 120:
                    rows.append([o1, o2])
                    i += 2
                    continue
        rows.append([options[i]])
        i += 1
    return rows


def render_question_to_cell(cell, q: Question, *, font_name: str, q_size: float, opt_size: float, ans_size: float, body_line_pts: float):
    # Clear default empty paragraph(s)
    for p in list(cell.paragraphs):
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass

    # Question line
    add_paragraph_with_text(
        cell,
        text=f"प्र.{q.no}. {q.question}".strip(),
        bold=False,
        font_name=font_name,
        font_size=q_size,
        line_pts=body_line_pts + 0.5,
        after=2,
    )

    # Options
    for row in layout_option_paragraphs(q.options):
        if len(row) == 2:
            # Pair in one paragraph with readable spacing
            o1, o2 = row
            paired = f"{o1.key} {o1.text}    {o2.key} {o2.text}"
            add_paragraph_with_text(
                cell,
                text=paired.strip(),
                bold=False,
                font_name=font_name,
                font_size=opt_size,
                line_pts=body_line_pts,
                after=1,
            )
        else:
            o = row[0]
            add_paragraph_with_text(
                cell,
                text=f"{o.key} {o.text}".strip(),
                bold=False,
                font_name=font_name,
                font_size=opt_size,
                line_pts=body_line_pts,
                after=1,
            )

    # Answer (if found)
    if q.correct:
        add_paragraph_with_text(
            cell,
            text=f"उत्तर: {q.correct}",
            bold=True,
            font_name=font_name,
            font_size=ans_size,
            line_pts=body_line_pts,
            after=2,
            before=2,
        )

    # Explanation + Facts (preserve both if present)
    combined_expl = q.explanation.strip()
    if q.facts.strip():
        if combined_expl:
            combined_expl = combined_expl + "\n" + "अन्य महत्वपूर्ण तथ्य: " + q.facts.strip()
        else:
            combined_expl = "अन्य महत्वपूर्ण तथ्य: " + q.facts.strip()

    if combined_expl:
        add_paragraph_with_text(
            cell,
            text=combined_expl,
            bold=False,
            font_name=font_name,
            font_size=ans_size,
            line_pts=body_line_pts,
            after=6,
        )


def build_output_doc(questions: List[Question], output_path: str, *, font_name: str):
    new_doc = Document()
    section = new_doc.sections[0]

    # 7 x 9 inch paper with 0.4 inch margins
    section.page_width = Inches(7)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)

    # Title
    p = new_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RBD PUBLICATION — Formatted Output")
    r.font.size = Pt(12)
    r.bold = True
    set_line_spacing(p, line_pts=16, space_after_pts=8)

    # Column widths
    content_w = 7 - 0.4 - 0.4
    div_w = 0.05
    col_w = (content_w - div_w) / 2

    COL_W = Inches(col_w)
    DIV_W = Inches(div_w)

    n = len(questions)
    half = (n + 1) // 2  # left column gets first half

    table = new_doc.add_table(rows=half, cols=3)
    table.autofit = False
    table.columns[0].width = COL_W
    table.columns[1].width = DIV_W
    table.columns[2].width = COL_W

    nil = {"val": "nil"}

    for i in range(half):
        row = table.rows[i]

        lc = row.cells[0]
        lc.width = COL_W
        set_cell_border(
            lc,
            top=nil,
            bottom=nil,
            left=nil,
            right={"val": "single", "sz": "4", "color": "AAAAAA", "space": "0"},
        )
        if i < len(questions):
            render_question_to_cell(
                lc,
                questions[i],
                font_name=font_name,
                q_size=8,
                opt_size=7.5,
                ans_size=7.5,
                body_line_pts=10,
            )

        mc = row.cells[1]
        mc.width = DIV_W
        set_cell_border(mc, top=nil, bottom=nil, left=nil, right=nil)

        rc = row.cells[2]
        rc.width = COL_W
        set_cell_border(rc, top=nil, bottom=nil, left=nil, right=nil)
        j = i + half
        if j < n:
            render_question_to_cell(
                rc,
                questions[j],
                font_name=font_name,
                q_size=8,
                opt_size=7.5,
                ans_size=7.5,
                body_line_pts=10,
            )

    new_doc.save(output_path)


def main():
    ap = argparse.ArgumentParser(description="Format MCQ DOCX into print-ready 2-column layout.")
    ap.add_argument("--input", required=True, help="Input .docx path")
    ap.add_argument("--output", default="Formatted_Output.docx", help="Output .docx path")
    ap.add_argument("--font", default="Mangal", help="Word font name (Hindi-friendly recommended)")
    args = ap.parse_args()

    doc = Document(args.input)
    # Keep paragraphs as-is; preserve internal line breaks inside paragraphs.
    paragraph_texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    blocks = split_questions_by_paragraphs(paragraph_texts)
    if not blocks:
        raise SystemExit("No question blocks detected. Check question markers like 'प्रश्न 1' or 'Q.1'.")

    questions: List[Question] = []
    for q_no, block_paras in blocks:
        q = parse_question_block(q_no, block_paras)
        questions.append(q)

    build_output_doc(questions, args.output, font_name=args.font)
    print(f"Parsed {len(questions)} questions. Saved: {args.output}")


if __name__ == "__main__":
    main()

