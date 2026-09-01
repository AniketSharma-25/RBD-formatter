# new august update # tagged approach trail   [ Tags required]

import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re
import tempfile
import os
import base64
import json
import subprocess
import shutil
import zipfile
import unicodedata
from reportlab.lib import colors
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PIL import Image as PILImage

# ================= AUTH SYSTEM =================
import psycopg2
import bcrypt
import datetime
import random
import smtplib
import uuid
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv

load_dotenv()


def _get_setting(key, default=""):
    """Pehle st.secrets (Streamlit Cloud) try karo, fir .env fallback (local dev)."""
    try:
        return st.secrets.get(key, os.getenv(key, default))
    except Exception:
        return os.getenv(key, default)


GMAIL_EMAIL = _get_setting("GMAIL_EMAIL")
GMAIL_APP_PASSWORD = _get_setting("GMAIL_APP_PASSWORD")
ADMIN_EMAIL = _get_setting("ADMIN_EMAIL")
DB_URL = _get_setting("DB_URL")  # Postgres connection string (Supabase/Neon/etc.)

ADMIN_EMAILS = {
    e.strip() for e in _get_setting("ADMIN_EMAIL", "").split(",") if e.strip()
}
GRANTED_USERS = {
    e.strip() for e in _get_setting("GRANTED_USERS", "").split(",") if e.strip()
}


def get_conn():
    return psycopg2.connect(DB_URL)


def init_db():
    conn = get_conn()
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        email TEXT PRIMARY KEY,
        created_at TEXT,
        is_admin BOOLEAN DEFAULT FALSE,
        can_format BOOLEAN DEFAULT FALSE,
        password_hash TEXT
    )''')
    # Purani DB (bina password_hash column ke) ke liye safe upgrade
    c.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS password_hash TEXT")
    c.execute('''CREATE TABLE IF NOT EXISTS otp_codes (
        email TEXT,
        code TEXT,
        expires_at TEXT
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS sessions (
        token TEXT PRIMARY KEY,
        email TEXT,
        created_at TEXT,
        expires_at TEXT,
        is_revoked BOOLEAN DEFAULT FALSE
    )''')
    conn.commit()
    conn.close()


def seed_trusted_accounts():
    """ADMIN_EMAIL aur GRANTED_USERS secrets se users ko auto-register/permission deta hai."""
    conn = get_conn()
    c = conn.cursor()
    now = datetime.datetime.now().isoformat()
    for email in ADMIN_EMAILS:
        c.execute(
            "INSERT INTO users (email, created_at, is_admin, can_format) VALUES (%s, %s, TRUE, TRUE) "
            "ON CONFLICT (email) DO UPDATE SET is_admin=TRUE, can_format=TRUE",
            (email, now)
        )
    for email in GRANTED_USERS:
        c.execute(
            "INSERT INTO users (email, created_at, is_admin, can_format) VALUES (%s, %s, FALSE, TRUE) "
            "ON CONFLICT (email) DO UPDATE SET can_format=TRUE",
            (email, now)
        )
    conn.commit()
    conn.close()


def add_user(email, is_admin=False):
    conn = get_conn()
    c = conn.cursor()
    try:
        now = datetime.datetime.now().isoformat()
        c.execute(
            "INSERT INTO users (email, created_at, is_admin, can_format) VALUES (%s, %s, %s, %s) "
            "ON CONFLICT (email) DO NOTHING",
            (email, now, is_admin, is_admin)
        )
        conn.commit()
    except Exception:
        pass
    conn.close()


def get_user(email):
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT email, is_admin, can_format FROM users WHERE email=%s", (email,))
    row = c.fetchone()
    conn.close()
    if row:
        return {"email": row[0], "is_admin": bool(row[1]), "can_format": bool(row[2])}
    return None


def generate_otp():
    return str(random.randint(100000, 999999))


def send_otp_email(email, code):
    try:
        msg = MIMEMultipart()
        msg["From"] = GMAIL_EMAIL
        msg["To"] = email
        msg["Subject"] = "OTP – RBD Formatter"
        msg.attach(MIMEText(f"Your OTP is: {code}\n\nValid for 10 minutes.", "plain"))
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(GMAIL_EMAIL, GMAIL_APP_PASSWORD)
        server.sendmail(GMAIL_EMAIL, email, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"SMTP ERROR: {e}")
        return False


def store_otp(email, code):
    conn = get_conn()
    c = conn.cursor()
    c.execute("DELETE FROM otp_codes WHERE email=%s", (email,))
    expiry = (datetime.datetime.now() + datetime.timedelta(minutes=10)).isoformat()
    c.execute("INSERT INTO otp_codes VALUES (%s, %s, %s)", (email, code, expiry))
    conn.commit()
    conn.close()


def verify_otp(email, code):
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT code, expires_at FROM otp_codes WHERE email=%s", (email,))
    row = c.fetchone()
    conn.close()
    if row and row[0] == code:
        if datetime.datetime.now() < datetime.datetime.fromisoformat(row[1]):
            return True
    return False


def create_session(email):
    token = str(uuid.uuid4())
    now = datetime.datetime.now()
    expires = now + datetime.timedelta(days=36500)  # ~100 saal = practically permanent
    conn = get_conn()
    c = conn.cursor()
    c.execute(
        "INSERT INTO sessions (token, email, created_at, expires_at, is_revoked) VALUES (%s, %s, %s, %s, FALSE)",
        (token, email, now.isoformat(), expires.isoformat())
    )
    conn.commit()
    conn.close()
    return token


def validate_session(token):
    if not token:
        return None
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT email, is_revoked FROM sessions WHERE token=%s", (token,))
    row = c.fetchone()
    conn.close()
    if not row:
        return None
    email, is_revoked = row
    if is_revoked:
        return None
    return get_user(email)


def revoke_user_sessions(email):
    conn = get_conn()
    conn.cursor().execute("UPDATE sessions SET is_revoked=TRUE WHERE email=%s", (email,))
    conn.commit()
    conn.close()


def revoke_session(token):
    if not token:
        return
    conn = get_conn()
    conn.cursor().execute("UPDATE sessions SET is_revoked=TRUE WHERE token=%s", (token,))
    conn.commit()
    conn.close()


def set_password(email, password):
    hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
    conn = get_conn()
    conn.cursor().execute("UPDATE users SET password_hash=%s WHERE email=%s", (hashed, email))
    conn.commit()
    conn.close()


def check_password(email, password):
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT password_hash FROM users WHERE email=%s", (email,))
    row = c.fetchone()
    conn.close()
    if not row or not row[0]:
        return False
    return bcrypt.checkpw(password.encode(), row[0].encode())


def has_password(email):
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT password_hash FROM users WHERE email=%s", (email,))
    row = c.fetchone()
    conn.close()
    return bool(row and row[0])


def _complete_login(user):
    token = create_session(user["email"])
    st.session_state["authenticated"] = True
    st.session_state["user_email"] = user["email"]
    st.session_state["is_admin"] = user["is_admin"]
    st.session_state["can_format"] = user["can_format"]
    st.session_state["session_token"] = token
    st.query_params["session"] = token


def login_page():
    st.title("🔐 Login")

    mode = st.radio(
        "Login mode",
        ["Password Login", "New here? Register / Forgot Password (OTP)"],
        horizontal=True,
    )

    # ---------------- PASSWORD LOGIN (returning users) ----------------
    if mode == "Password Login":
        email = st.text_input("Email", key="pw_email")
        password = st.text_input("Password", type="password", key="pw_pass")
        if st.button("Login"):
            if email and password and check_password(email, password):
                user = get_user(email)
                _complete_login(user)
                st.success("Logged in!")
                st.rerun()
            else:
                st.error("Invalid email or password")
        st.caption("Naya user ho ya password bhool gaye ho? Doosra option chuno — OTP se verify karke password set kar sakte ho.")

    # ---------------- OTP REGISTRATION / PASSWORD RESET ----------------
    else:
        email = st.text_input("Email", key="otp_email_input")
        if st.button("Send OTP"):
            if email:
                user = get_user(email)
                if not user:
                    add_user(email, email in ADMIN_EMAILS)
                otp = generate_otp()
                sent = send_otp_email(email, otp)
                if sent:
                    store_otp(email, otp)
                    st.session_state["otp_email"] = email
                    st.session_state.pop("otp_verified_email", None)
                    st.success("OTP sent to your email ✅")
                else:
                    st.error("❌ Failed to send OTP email")

        if "otp_email" in st.session_state and "otp_verified_email" not in st.session_state:
            code = st.text_input("Enter OTP", type="password", key="otp_code_input")
            if st.button("Verify OTP"):
                if verify_otp(st.session_state["otp_email"], code):
                    st.session_state["otp_verified_email"] = st.session_state["otp_email"]
                    st.success("OTP verified ✅ — ab neeche apna password set karo")
                    st.rerun()
                else:
                    st.error("Invalid OTP")

        if "otp_verified_email" in st.session_state:
            new_pass = st.text_input("Set a password", type="password", key="new_pass_1")
            confirm_pass = st.text_input("Confirm password", type="password", key="new_pass_2")
            if st.button("Set Password & Login"):
                if len(new_pass) < 6:
                    st.error("Password kam se kam 6 characters ka rakho")
                elif new_pass != confirm_pass:
                    st.error("Passwords match nahi kar rahe")
                else:
                    verified_email = st.session_state["otp_verified_email"]
                    set_password(verified_email, new_pass)
                    user = get_user(verified_email)
                    _complete_login(user)
                    st.session_state.pop("otp_email", None)
                    st.session_state.pop("otp_verified_email", None)
                    st.success("Password set! Logged in!")
                    st.rerun()

    if not st.session_state.get("authenticated"):
        st.stop()


# =============================================================================
# TEXT UTILITIES  (unchanged)
# =============================================================================
def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'\(.*?\d{2}.*?\[.*?\].*?\(.*?\).*?\)', '', text)
    text = re.sub(r'(?<![^\s।])प्रश्न\s+\d+\s*', '', text)
    text = re.sub(r'^\d+\.\s*', '', text)
    text = re.sub(r'^\.+\s*', '', text)
    text = text.replace('\t', ' ')
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


# =============================================================================
# TABLE EDITOR — unchanged from original
# =============================================================================
import re as _re

def split_row_by_char(text: str, char: str) -> list:
    if not char:
        return [text.strip()]
    escaped = _re.escape(char)
    parts = _re.split(escaped, text)
    return [p.strip() for p in parts]


def auto_split_row(text: str) -> list:
    if "|" in text:
        return [p.strip() for p in text.split("|") if p.strip()]
    if "\t" in text:
        return [p.strip() for p in text.split("\t")]
    if _re.search(r"\s{2,}", text):
        return [p.strip() for p in _re.split(r"\s{2,}", text)]
    return [text.strip()]


def merge_lines_to_table(lines: list, split_char: str) -> list:
    rows = []
    for line in lines:
        if split_char == "__auto__":
            cols = auto_split_row(line)
        else:
            cols = split_row_by_char(line, split_char)
        rows.append(cols)
    max_cols = max(len(r) for r in rows) if rows else 1
    for row in rows:
        while len(row) < max_cols:
            row.append("")
    return rows


def _decompose_question_to_lines(q: dict) -> list:
    lines = []
    if q.get("question"):
        for raw_line in q["question"].split("\n"):
            raw_line = raw_line.strip()
            if raw_line and raw_line != "[SUCHI_HEADER]":
                lines.append(("question_body", raw_line))
    for line in q.get("kathan_lines", []):
        lines.append(("kathan_line", line))
    for left, right in q.get("suchi_rows", []):
        lines.append(("suchi_row", f"{left}  |  {right}"))
    for left, right in q.get("match_following_rows", []):
        lines.append(("match_row", f"{left}  |  {right}"))
    for opt in q.get("options", []):
        lines.append(("option", f"{opt['key']} {opt['text']}"))
    if q.get("explanation"):
        for raw_line in q["explanation"].split("|"):
            raw_line = raw_line.strip()
            if raw_line:
                lines.append(("explanation", raw_line))
    return lines


def _apply_table_to_question(questions_ref, q_idx, rows, dest_key, h1=None, h2=None, selected_lines=None, split_char="–"):
    q = questions_ref[q_idx]
    pairs = []
    for row in rows:
        left  = row[0] if len(row) > 0 else ""
        right = row[1] if len(row) > 1 else ""
        if len(row) > 2:
            right = "  ".join(row[1:])
        pairs.append((left, right))
    q[dest_key] = pairs
    if dest_key == "suchi_rows" and h1 and h2:
        q["suchi_col_headers"] = (h1, h2)
    st.session_state.setdefault("te_log", []).append({
        "q_no":       q["no"],
        "dest":       dest_key,
        "rows":       len(pairs),
        "cols":       max(2, max(len(r) for r in rows)),
        "split_char": split_char,
    })


def _init_table_editor_state(questions):
    if "te_questions" not in st.session_state or st.session_state.get("te_source") != id(questions):
        import copy
        st.session_state["te_questions"] = copy.deepcopy(questions)
        st.session_state["te_source"]    = id(questions)
        st.session_state["te_log"]       = []


def render_table_editor_tab(questions_ref):
    st.markdown("### 🔧 Manual Table Editor")
    st.caption(
        "Use this when the parser missed a suchi / match-the-following table. "
        "Select the raw lines from a question, set a split character, and apply — "
        "they become `match_following_rows` rendered as a 2-col table in the DOCX."
    )
    if not questions_ref:
        st.info("No questions parsed yet.")
        return

    real_indices = [i for i, q in enumerate(questions_ref) if not q.get('is_separator')]
    if not real_indices:
        st.info("No questions parsed yet.")
        return
    q_labels = [f"Q{questions_ref[i]['no']} — {questions_ref[i]['question'][:60]}…" for i in real_indices]
    sel      = st.selectbox("Select question to inspect / fix", range(len(q_labels)),
                            format_func=lambda i: q_labels[i])
    q_idx = real_indices[sel]
    q = questions_ref[q_idx]
    st.markdown("---")

    raw_lines = _decompose_question_to_lines(q)
    if not raw_lines:
        st.warning("No raw lines found for this question.")
        return

    st.markdown(f"**Raw lines for Q{q['no']}** — tick the lines you want to convert to a table:")
    selected_indices = []
    for i, (key, text) in enumerate(raw_lines):
        label = f"`[{key}]`  {text[:100]}"
        if st.checkbox(label, key=f"te_line_{q_idx}_{i}"):
            selected_indices.append(i)

    if not selected_indices:
        st.info("☝️ Tick 2 or more lines above to enable the table builder.")
        return

    selected_lines = [raw_lines[i][1] for i in selected_indices]
    st.markdown("---")

    col_mode, col_char = st.columns([2, 1])
    with col_mode:
        mode = st.radio(
            "Split mode",
            ["User Formatted Table (specify character)", "Default Table (auto-detect)"],
            horizontal=True,
            key=f"te_mode_{q_idx}",
        )
    with col_char:
        split_char = "__auto__"
        if "User Formatted" in mode:
            split_char = st.text_input(
                "Split character", value="–",
                max_chars=5, key=f"te_char_{q_idx}",
                help="e.g.  –  |  -  :  →"
            )

    preview_rows = merge_lines_to_table(selected_lines, split_char)
    max_cols     = max(len(r) for r in preview_rows)

    st.markdown("**Live preview:**")
    header_html = "".join(
        f"<th style='padding:4px 10px;background:#263238;color:#9db4c0;border:1px solid #333;'>Col {ci+1}</th>"
        for ci in range(max_cols)
    )
    rows_html = ""
    for ri, row in enumerate(preview_rows):
        bg = "#1a1d27" if ri % 2 == 0 else "#1e2232"
        cells = "".join(
            f"<td style='padding:4px 10px;border:1px solid #333;color:#dde1f0;'>{cell}</td>"
            for cell in row
        )
        rows_html += f"<tr style='background:{bg};'>{cells}</tr>"

    st.markdown(
        f"""<table style='border-collapse:collapse;font-size:13px;width:100%;'>
              <thead><tr>{header_html}</tr></thead>
              <tbody>{rows_html}</tbody>
            </table>""",
        unsafe_allow_html=True,
    )
    st.caption(f"{len(preview_rows)} rows × {max_cols} cols")
    st.markdown("---")

    dest = st.radio(
        "Store result as",
        ["match_following_rows  (renders as match-the-following table)",
         "suchi_rows            (renders as सूची-I / सूची-II table)"],
        key=f"te_dest_{q_idx}",
    )

    if "suchi" in dest:
        col_h1, col_h2 = st.columns(2)
        with col_h1:
            h1 = st.text_input("Left header",  value="सूची-I",  key=f"te_h1_{q_idx}")
        with col_h2:
            h2 = st.text_input("Right header", value="सूची-II", key=f"te_h2_{q_idx}")
    else:
        h1 = h2 = None

    if st.button("✅ Apply as Table", type="primary", key=f"te_apply_{q_idx}"):
        _apply_table_to_question(
            questions_ref, q_idx, preview_rows,
            dest_key="suchi_rows" if "suchi" in dest else "match_following_rows",
            h1=h1, h2=h2,
            selected_lines=selected_lines,
            split_char=split_char,
        )
        st.success(
            f"✅ Applied! Q{q['no']} now has "
            f"{'suchi_rows' if 'suchi' in dest else 'match_following_rows'} "
            f"with {len(preview_rows)} rows × {max_cols} cols. "
            "Re-generate DOCX to see the result."
        )
        st.rerun()

    log = st.session_state.get("te_log", [])
    if log:
        st.markdown("---")
        st.markdown("**Operation log:**")
        for entry in reversed(log[-5:]):
            st.markdown(
                f"✓ Q{entry['q_no']} · {entry['dest']} · "
                f"{entry['rows']}r × {entry['cols']}c · split=`{entry['split_char']}`"
            )


# =============================================================================
# FONT CONFIGURATION  (unchanged)
# =============================================================================
HINDI_FONTS = {
    "Kokila": "Kokila",
    "Mangal": "Mangal",
    "Nirmala UI": "Nirmala UI",
    "Aparajita": "Aparajita",
    "Utsaah": "Utsaah",
    "Kruti Dev 010": "Kruti Dev 010",
    "Devanagari New": "Devanagari New",
}
ENGLISH_FONTS = {
    "Arial": "Arial",
    "Times New Roman": "Times New Roman",
    "Calibri": "Calibri",
    "Georgia": "Georgia",
    "Cambria": "Cambria",
    "Garamond": "Garamond",
    "Trebuchet MS": "Trebuchet MS",
    "Verdana": "Verdana",
    "Book Antiqua": "Book Antiqua",
    "Century Gothic": "Century Gothic",
}

# =============================================================================
# PAGE CONFIG & INIT
# =============================================================================
st.set_page_config(page_title="RBD Formatter", layout="wide")
st.title("📚 RBD Publication – Smart Formatter")
init_db()

if not st.session_state.get("_seeded"):
    seed_trusted_accounts()
    st.session_state["_seeded"] = True

if not st.session_state.get("authenticated"):
    token = st.query_params.get("session")
    if token:
        user = validate_session(token)
        if user:
            st.session_state["authenticated"] = True
            st.session_state["user_email"] = user["email"]
            st.session_state["is_admin"] = user["is_admin"]
            st.session_state["can_format"] = user["can_format"]
            st.session_state["session_token"] = token
        else:
            st.query_params.clear()

if not st.session_state.get("authenticated"):
    login_page()

if st.session_state.get("is_admin"):
    st.sidebar.title("👑 Admin Panel")
    conn = get_conn()
    c = conn.cursor()
    c.execute("SELECT email, can_format FROM users")
    users = c.fetchall()
    conn.close()
    for email, can_format in users:
        col1, col2 = st.sidebar.columns([3, 1])
        with col1:
            val = st.checkbox(email, value=bool(can_format), key=f"perm_{email}")
        with col2:
            if st.sidebar.button("🚫", key=f"revoke_{email}", help=f"Revoke sessions for {email}"):
                revoke_user_sessions(email)
                st.sidebar.success(f"Sessions revoked for {email}")
                st.rerun()
        if val != bool(can_format):
            conn = get_conn()
            conn.cursor().execute("UPDATE users SET can_format=%s WHERE email=%s", (val, email))
            conn.commit()
            conn.close()
            st.rerun()

if st.session_state.get("authenticated"):
    st.sidebar.markdown("---")
    st.sidebar.write(f"👤 {st.session_state.get('user_email', '')}")
    if st.sidebar.button("🚪 Logout"):
        revoke_session(st.session_state.get("session_token"))
        st.session_state.clear()
        st.query_params.clear()
        st.rerun()

if not st.session_state.get("can_format"):
    st.error("❌ You are not allowed to use formatter")
    st.stop()

# =============================================================================
# FILE UPLOAD — single or multiple (max 25), reorderable, mixed .docx/.txt/.json
# =============================================================================
def _file_key(f):
    return f"{f.name}::{f.size}"


st.subheader("📄 Upload Question Paper(s)")
upload_mode = st.radio(
    "Upload mode",
    [
        "Single file",
        "Multiple files (merge into one paper)",
        "Batch (process each file separately)",
    ],
    horizontal=True, key="upload_mode"
)

if upload_mode == "Single file":
    _single = st.file_uploader(
        "Upload a .docx, .txt, or .json file", type=["docx", "txt", "json"],
        accept_multiple_files=False, key="single_uploader"
    )
    raw_uploaded_files = [_single] if _single else []
else:
    raw_uploaded_files = st.file_uploader(
        "Upload up to 25 files (.docx, .txt, .json — any mix)", type=["docx", "txt", "json"],
        accept_multiple_files=True, key="multi_uploader"
    ) or []
    if len(raw_uploaded_files) > 25:
        st.error(f"⚠️ {len(raw_uploaded_files)} files uploaded — only the first 25 will be used.")
        raw_uploaded_files = raw_uploaded_files[:25]

add_filename_separator = False
uploaded_files_ordered = []
batch_output_names     = {}   # file_key -> custom output filename (no extension)

if raw_uploaded_files:
    keys_now = [_file_key(f) for f in raw_uploaded_files]
    file_by_key = {_file_key(f): f for f in raw_uploaded_files}

    if st.session_state.get("_file_order_set") != set(keys_now):
        st.session_state["_file_order"]     = keys_now
        st.session_state["_file_order_set"] = set(keys_now)

    order = [k for k in st.session_state.get("_file_order", keys_now) if k in file_by_key]
    for k in keys_now:
        if k not in order:
            order.append(k)
    st.session_state["_file_order"] = order

    if upload_mode == "Multiple files (merge into one paper)" and len(raw_uploaded_files) > 1:
        st.markdown("**📑 Adjust paper sequence** (papers are merged in this order — one ends, the next begins right after):")
        for i, k in enumerate(order):
            f = file_by_key[k]
            c1, c2, c3 = st.columns([8, 1, 1])
            c1.write(f"{i + 1}. {f.name}")
            if c2.button("⬆️", key=f"up_{k}", disabled=(i == 0)):
                order[i - 1], order[i] = order[i], order[i - 1]
                st.session_state["_file_order"] = order
                st.rerun()
            if c3.button("⬇️", key=f"dn_{k}", disabled=(i == len(order) - 1)):
                order[i + 1], order[i] = order[i], order[i + 1]
                st.session_state["_file_order"] = order
                st.rerun()

        add_filename_separator = st.checkbox(
            "Insert a bold filename heading between merged papers", value=False,
            key="add_fname_sep",
            help="Adds a bold row with the source filename right before each new paper begins."
        )

    elif upload_mode == "Batch (process each file separately)" and raw_uploaded_files:
        st.markdown("**📝 Each file is processed independently — set an output name for each:**")
        for k in order:
            f = file_by_key[k]
            default_name = f.name.rsplit('.', 1)[0]
            custom = st.text_input(
                f"Output name for **{f.name}**", value=default_name, key=f"batch_name_{k}"
            )
            batch_output_names[k] = custom

    uploaded_files_ordered = [file_by_key[k] for k in order]

# =============================================================================
# SIDEBAR – all settings  (unchanged)
# =============================================================================
with st.sidebar:
    st.header("📄 Page Design")
    page_width      = st.number_input("Page Width (inches)",    5.0, 12.0, 8.27,  0.1)
    page_height     = st.number_input("Page Height (inches)",   6.0, 14.0, 11.69,  0.1)
    top_margin      = st.number_input("Top Margin (inches)",    0.2,  1.0, 1.0,  0.05)
    bottom_margin   = st.number_input("Bottom Margin (inches)", 0.2,  1.0, 0.75, 0.05)
    left_margin     = st.number_input("Left Margin (inches)",   0.2,  1.0, 0.75, 0.05)
    right_margin    = st.number_input("Right Margin (inches)",  0.2,  1.0, 0.75, 0.05)

    st.header("📐 Layout")
    num_columns = st.selectbox("Number of Columns", [2, 3], index=0)
    auto_fill   = st.checkbox("Auto‑fill pages", True)

    st.header("🔢 Question Numbering")
    numbering_choice = st.radio(
        "Numbering mode",
        [
            "Sequenced — renumber 1…N across all papers (default)",
            "Default — keep the original numbers from each file",
            "Customized — start numbering from a chosen number",
        ],
        index=0,
        key="numbering_choice",
    )
    numbering_start = 1
    if numbering_choice.startswith("Customized"):
        numbering_start = st.number_input(
            "Start numbering from", min_value=1, value=1, step=1, key="numbering_start"
        )
    if numbering_choice.startswith("Sequenced"):
        numbering_mode = "sequenced"
    elif numbering_choice.startswith("Default"):
        numbering_mode = "default"
    else:
        numbering_mode = "customized"

    st.header("🔤 Font Settings (Output DOCX)")
    _default_font_lang_idx = 1 if st.session_state.get("_is_english_doc") else 0
    font_language = st.selectbox(
        "Select Font Language",
        ["Hindi (Devanagari)", "English"],
        index=_default_font_lang_idx
    )
    if font_language == "Hindi (Devanagari)":
        selected_font_name = st.selectbox("Select Hindi Font", list(HINDI_FONTS.keys()), index=0)
        FONT_DOCX = HINDI_FONTS[selected_font_name]
    else:
        selected_font_name = st.selectbox("Select English Font", list(ENGLISH_FONTS.keys()), index=0)
        FONT_DOCX = ENGLISH_FONTS[selected_font_name]
    st.caption(f"✅ Selected font: **{FONT_DOCX}** — applied to all text in output DOCX")

    st.markdown("---")
    st.header("🎨 Style Overrides")

    with st.expander("✍️ Question Style", expanded=True):
        q_font   = st.slider("Font size (pt)", 5.0, 15.0, 12.0, 0.5, key="q_font")
        q_bold   = st.checkbox("Bold question text", value=True, key="q_bold")
        st.markdown("**Indent levels**")
        level1_indent = st.number_input("Level-1 indent (in)", 0.0, 0.5, 0.0, 0.05, key="q_l1")
        level2_indent = st.number_input("Level-2 indent (in)", 0.05, 1.0, 0.15, 0.05, key="q_l2")
        q_indent = level2_indent
        line_spacing = st.slider("Line spacing (pt)",          8.0, 15.0,  9.5, 0.5, key="q_line")
        para_spacing = st.slider("Space after paragraph (pt)", 0.0,  6.0,  0.0, 0.5, key="q_para")
        char_spacing = st.slider("Character spacing (pt)",     0.0,  3.0,  0.0, 0.1, key="q_char")

    with st.expander("📋 Options Style", expanded=False):
        opt_font = st.slider("Font size (pt)", 5.0, 15.0, 12.0, 0.5, key="o_font")
        opt_bold = st.checkbox("Bold options", value=False, key="o_bold")
        opts_per_line = st.selectbox("Max options per line", [2, 3, 4], index=0, key="o_perline")
        if opts_per_line == 4:
            _default_char = 80
        elif opts_per_line == 3:
            _default_char = 68
        else:
            _default_char = 68
        opt_char_limit = st.slider(
            "Option line length threshold (chars)", 40, 120, _default_char, key="o_charlim"
        )
        show_correct_inline = st.checkbox(
            "Show correct answer on last option line (right-aligned)", True, key="o_showans"
        )

    with st.expander("📝 Explanation Style", expanded=False):
        expl_font   = st.slider("Font size (pt)", 5.0, 15.0, 12.0, 0.5, key="e_font")
        expl_bullet = st.checkbox("Bullet (➤) before व्याख्या heading", True, key="e_bullet")
        expl_bg     = st.checkbox("Light grey background for explanation", True, key="e_bg")
        _default_expl_label = "Explanation" if st.session_state.get("_is_english_doc") else "व्याख्या"
        expl_prefix = st.text_input(
            "Explanation label", value=_default_expl_label,
            help="The heading word shown before explanation text.",
            key="e_prefix"
        )

    with st.expander("📋 Kathan Style", expanded=False):
        kathan_font   = st.slider("Kathan font size (pt)", 5.0, 15.0, 11.5, 0.5, key="k_font")
        kathan_bold   = st.checkbox("Bold kathan lines", value=False, key="k_bold")
        kathan_indent_extra = st.number_input(
            "Extra indent for kathan lines (in)", 0.0, 0.5, 0.1, 0.05, key="k_indent",
            help="Added on top of level2_indent for kathan/statement lines"
        )
        kathan_bg     = st.checkbox("Light background for kathan block", value=False, key="k_bg")

    with st.expander("🔗 Match the Following Table", expanded=False):
        mtf_sep_choice = st.selectbox(
            "Separator type",
            ["Auto (any dash)", "Single dash  -", "Double dash  --", "En dash  –", "Arrow  →", "Tab"],
            index=0,
            key="mtf_sep"
        )
        _sep_map = {
            "Auto (any dash)":   None,
            "Single dash  -":    r'^(\([^)]+\)\s*.*?)\s+-\s+(\([^)]+\)\s*.*)$',
            "Double dash  --":   r'^(\([^)]+\)\s*.*?)\s+--\s+(\([^)]+\)\s*.*)$',
            "En dash  –":        r'^(\([^)]+\)\s*.*?)\s+–\s+(\([^)]+\)\s*.*)$',
            "Arrow  →":          r'^(\([^)]+\)\s*.*?)\s+→\s+(\([^)]+\)\s*.*)$',
            "Tab":               r'^(\([^)]+\)\s*.*?)\t(\([^)]+\)\s*.*)$',
        }
        mtf_sep_pattern = _sep_map[mtf_sep_choice]

        mtf_col1, mtf_col2 = st.columns(2)
        with mtf_col1:
            mtf_h1 = st.text_input("Left header",  value="सूची-I",  key="mtf_h1")
        with mtf_col2:
            mtf_h2 = st.text_input("Right header", value="सूची-II", key="mtf_h2")

        mtf_left_pct = st.slider(
            "Left column width (%)", 30, 70, 50, 5, key="mtf_pct"
        )
        mtf_show_header = st.checkbox("Show shaded header row", value=False, key="mtf_hdr")
        mtf_inside_v    = st.checkbox("Show border between columns", value=False, key="mtf_vbdr")
        mtf_cell_pad    = st.number_input("Cell left padding (twips)", 0, 120, 40, 10, key="mtf_pad")

    st.header("📝 Header & Footer")
    chapter_heading = st.text_input("Chapter Heading", "")
    header_font  = st.slider("Header font size (pt)", 8.0, 16.0, 13.0, 0.5)
    header_bold  = st.checkbox("Header bold", True)
    header_bg    = st.checkbox("Header grey background", True)
    header_align = st.selectbox("Header alignment", ["Left", "Center", "Right"], index=1)
    header_top_padding = st.slider("Header top padding (twips)", 0, 400, 60, 20)

    book_name  = "RBD PUBLICATION"
    topic_name = ""

    st.header("🔢 Page Numbers")
    page_num_pos = st.selectbox(
        "Position",
        ["None","Top Left","Top Center","Top Right","Bottom Left","Bottom Center","Bottom Right"],
        index=5
    )
    hide_on_first = st.checkbox("Hide on first page", False) if page_num_pos != "None" else False

    st.header("✨ Extras")
    show_separator   = st.checkbox("Show line after each question", False)
    include_metadata = st.checkbox("Include PYQ metadata in output", False)

    if st.checkbox("Extra compact mode", False):
        line_spacing = 5.0
        para_spacing = 0.0
        q_font       = 5.0
        opt_font     = 5.0
        expl_font    = 5.0

# =============================================================================
# ███████████████████████████████████████████████████████████████████████████
# NEW TAG-BASED PARSER  (replaces parse_questions + parse_english_questions)
# ███████████████████████████████████████████████████████████████████████████
# =============================================================================

# ---------------------------------------------------------------------------
# STEP 1 — Extract full text from all paragraphs (preserve newlines)
# ---------------------------------------------------------------------------
def _doc_to_text(doc):
    """Return the full document as one string, paragraphs separated by newlines."""
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# STEP 2 — Split document text into tag blocks
# ---------------------------------------------------------------------------
# Supported tags: question, kathan, suchi, option, answer, expl
_TAG_RE = re.compile(
    r'<(question|kathan|suchi|option|answer|expl)>(.*?)</\1>',
    re.DOTALL | re.IGNORECASE
)

def _split_into_tag_blocks(text):
    """
    Returns list of (tag_name, content) pairs in document order.
    Content has leading/trailing whitespace stripped.
    Falls back to raw-text heuristic if no tags found at all.
    """
    blocks = [(m.group(1).lower(), m.group(2).strip())
              for m in _TAG_RE.finditer(text)]
    return blocks


# ---------------------------------------------------------------------------
# STEP 3 — Group tag blocks into per-question dicts
# ---------------------------------------------------------------------------
def _group_blocks_into_questions(blocks):
    """
    A new question starts whenever we see a <question> or <kathan> tag
    that contains a question number.
    Returns list of raw question dicts:
      {no, question_raw, kathan_raw, suchi_raw, option_raw, answer_raw, expl_raw}
    """
    questions = []
    current   = None

    def _new_q():
        return {
            "no":           "",
            "question_raw": "",
            "kathan_raw":   "",
            "suchi_raw":    "",
            "option_raw":   "",
            "answer_raw":   "",
            "expl_raw":     "",
        }

    for tag, content in blocks:
        if tag == "question":
            # Extract question number from content
            m = re.match(r'^\s*(?:प्रश्न\s*)?(\d+)[.):\s]', content)
            if not m:
                m = re.match(r'^\s*(?:Q|q)\.?\s*(\d+)[.):\s]?', content)
            if m:
                if current:
                    questions.append(current)
                current = _new_q()
                current["no"]           = m.group(1)
                current["question_raw"] = content
            else:
                # No number — append to last question body
                if current:
                    current["question_raw"] += "\n" + content

        elif tag == "kathan":
            if current is None:
                current = _new_q()
            current["kathan_raw"] += ("\n" if current["kathan_raw"] else "") + content

        elif tag == "suchi":
            if current is None:
                current = _new_q()
            current["suchi_raw"] += ("\n" if current["suchi_raw"] else "") + content

        elif tag == "option":
            if current is None:
                current = _new_q()
            current["option_raw"] += ("\n" if current["option_raw"] else "") + content

        elif tag == "answer":
            if current is None:
                current = _new_q()
            current["answer_raw"] = content

        elif tag == "expl":
            if current is None:
                current = _new_q()
            current["expl_raw"] += ("\n" if current["expl_raw"] else "") + content

    if current:
        questions.append(current)

    return questions


# ---------------------------------------------------------------------------
# STEP 4 — Parse <suchi> content: handle all 96 variants
# Separators: │  |  -  --   (pipe variants and dash variants)
# Left markers:  (A-D)  (1-4)  (अ-द)  (i-iv)  (I-IV) and bare A. 1. i. अ. etc.
# Right markers: same set as left
# Header line: "सूची-I (subtitle)  सूची-II (subtitle)"  or "सूची – A  सूची – B"
# ---------------------------------------------------------------------------

# Pattern to detect a suchi separator in a line
_SUCHI_SEP_RE = re.compile(
    r'(?:│|\|{1,2}|--|-(?!-)|–)',  # │, |, ||, --, –  but not part of text
)

# Pattern for the LEFT item of a suchi row:
# Matches: (A) text,  (1) text,  (अ) text,  (i) text,  A. text,  1. text, etc.
_SUCHI_LEFT_ITEM_RE = re.compile(
    r'^\s*'
    r'(?:'
    r'\(([A-Da-d1-4ivxIVXअबसदabcd])\)'   # (A), (1), (i), (अ) style
    r'|([A-Da-d1-4ivxIVX])\.'              # A. 1. i. bare dot style
    r'|([अबसद])\.'                          # अ. ब. स. द. bare dot style
    r')\s*(.+)',
    re.UNICODE
)

def _detect_suchi_header(text):
    """
    Detect and return (h1_label, h2_label, remaining_text) from a suchi block.
    Handles both  सूची-I / सूची-II  and  सूची – A / सूची – B  styles.
    Also handles "List-I / List-II" for English.
    """
    # Normalize the text — collapse multiple spaces
    text = re.sub(r'[ \t]+', ' ', text)

    # Pattern 1: सूची-I (sub1)  सूची-II (sub2)   — hyphen style
    m = re.search(
        r'(सूची[-\s–]I(?:\s*\([^)]*\))?)'
        r'\s+'
        r'(सूची[-\s–]II(?:\s*\([^)]*\))?)',
        text, re.UNICODE
    )
    if m:
        h1 = m.group(1).strip()
        h2 = m.group(2).strip()
        remaining = text[m.end():].strip()
        return h1, h2, remaining

    # Pattern 2: सूची – A  सूची – B   — en-dash style
    m = re.search(
        r'(सूची\s*–\s*[AB](?:\s*\([^)]*\))?)'
        r'\s+'
        r'(सूची\s*–\s*[AB](?:\s*\([^)]*\))?)',
        text, re.UNICODE
    )
    if m:
        h1 = m.group(1).strip()
        h2 = m.group(2).strip()
        remaining = text[m.end():].strip()
        return h1, h2, remaining

    # Pattern 3: List-I  List-II  (English)
    m = re.search(
        r'(List[-\s]I(?:\s*\([^)]*\))?)'
        r'\s+'
        r'(List[-\s]II(?:\s*\([^)]*\))?)',
        text, re.IGNORECASE
    )
    if m:
        h1 = m.group(1).strip()
        h2 = m.group(2).strip()
        remaining = text[m.end():].strip()
        return h1, h2, remaining

    # No header found — use defaults
    return "सूची-I", "सूची-II", text


def _parse_suchi_rows(text):
    """
    Parse suchi data rows from text. Returns list of (left, right) string pairs.
    Handles all separator types: │ | - -- –
    Handles all marker styles: (A-D), (1-4), (अ-द), (i-iv), A., 1., etc.

    Strategy:
    1. Split text into lines.
    2. For each non-empty line that contains a suchi separator, split on it.
    3. Clean up each side.
    """
    rows = []
    lines = [l.strip() for l in text.replace('\r', '\n').split('\n') if l.strip()]

    for line in lines:
        # Skip pure header lines (lines that contain सूची but no separator items)
        if re.search(r'सूची|List-I|List-II', line, re.IGNORECASE):
            has_item = _SUCHI_LEFT_ITEM_RE.match(line)
            if not has_item:
                continue  # it's a header line, skip

        # Try to split on a separator
        # Priority: │ (box char) > || > -- > | > - > –
        sep_found = None
        for sep_pattern in [r'│', r'\|\|', r'--', r'\|', r'(?<!\-)-(?!\-)', r'–']:
            m = re.search(sep_pattern, line)
            if m:
                left_part  = line[:m.start()].strip()
                right_part = line[m.end():].strip()
                if left_part and right_part:
                    sep_found = (left_part, right_part)
                    break

        if sep_found:
            rows.append(sep_found)

    return rows


def parse_suchi_tagged(suchi_raw):
    """
    Master suchi parser for <suchi>…</suchi> content.
    Returns (h1, h2, rows) where rows = [(left_str, right_str), ...]
    Handles all 96 documented variants.
    """
    if not suchi_raw:
        return "सूची-I", "सूची-II", []

    h1, h2, body = _detect_suchi_header(suchi_raw)
    rows = _parse_suchi_rows(body)

    # Fallback: if no rows found but body has content, try the whole text including header region
    if not rows:
        rows = _parse_suchi_rows(suchi_raw)

    return h1, h2, rows


# ---------------------------------------------------------------------------
# STEP 5 — Parse <kathan> content
# Handles:
#   कथन (A): …  / कारण (R): …      → Assertion-Reason style
#   a. … / b. … / c. …             → Statement list style
#   1. … / 2. … / 3. …             → Numbered statement style
#   Statement 1: … / Statement 2:  → English statement style
# Returns list of plain strings (one per line), ready for indented rendering
# ---------------------------------------------------------------------------

def parse_kathan_lines(kathan_raw):
    """
    Returns list of strings. Each string is one kathan/statement line.
    Preserves the label prefix (कथन (A):, a., 1., Statement 1:, etc.)
    """
    if not kathan_raw:
        return []

    lines = []
    # Split on newlines first
    raw_lines = [l.strip() for l in kathan_raw.replace('\r', '\n').split('\n') if l.strip()]

    # If it's all on one line (common in tagged format), try splitting on known patterns
    if len(raw_lines) == 1:
        text = raw_lines[0]
        # Split on:  कथन/कारण prefix, or  a./b./c. prefix, or  1./2./3. prefix
        # Insert newline before each one
        text = re.sub(
            r'(?<=[।\.\?])\s+'
            r'(?='
            r'(?:कथन|कारण|Statement|Reason|Assertion)'
            r'|[a-eA-E]\.'
            r'|\d\.'
            r'|\([a-eA-E]\)'
            r'|\(\d\)'
            r')',
            '\n', text
        )
        raw_lines = [l.strip() for l in text.split('\n') if l.strip()]

    for line in raw_lines:
        # Normalize common line-start patterns
        # कथन (A): text  →  keep as is
        # कारण (R): text  →  keep as is
        # a. text  →  keep as is
        # (a) text  →  keep as is
        if line:
            lines.append(line)

    return lines


# ---------------------------------------------------------------------------
# STEP 6 — Parse <option> content
# Handles: (a) text (b) text … on same line OR separate lines
# Also handles  कूट:  prefix line
# ---------------------------------------------------------------------------

_OPT_RE = re.compile(r'\(([a-dA-D])\)\s*(.*?)(?=\([a-dA-D]\)|$)', re.DOTALL)
_OPT_LINE_RE = re.compile(r'^\s*\(([a-dA-D])\)\s*(.+)$')

def parse_options_tagged(option_raw):
    """
    Returns list of {key, text} dicts.
    Handles options on one line OR one-per-line (with or without कूट: header).
    """
    if not option_raw:
        return []

    # Remove कूट: header line if present
    text = re.sub(r'^\s*(?:कूट|Codes?)\s*[:–-]?\s*\n?', '', option_raw, flags=re.IGNORECASE).strip()

    options = []

    # Try line-by-line first (most reliable)
    for line in text.replace('\r', '\n').split('\n'):
        line = line.strip()
        if not line:
            continue
        m = _OPT_LINE_RE.match(line)
        if m:
            options.append({"key": f"({m.group(1).lower()})", "text": m.group(2).strip()})

    # If nothing found, try inline parsing (all options on one line)
    if not options:
        for m in _OPT_RE.finditer(text):
            opt_text = re.sub(r'\s+', ' ', m.group(2)).strip()
            if opt_text:
                options.append({"key": f"({m.group(1).lower()})", "text": opt_text})

    return options[:4]


# ---------------------------------------------------------------------------
# STEP 7 — Parse <answer> and <expl>
# ---------------------------------------------------------------------------

def parse_answer_tagged(answer_raw):
    """Returns correct answer string like '(b)' or empty string."""
    if not answer_raw:
        return ""
    m = re.search(r'\(([a-dA-D])\)', answer_raw)
    return f"({m.group(1).lower()})" if m else ""


def parse_expl_tagged(expl_raw):
    """Returns clean explanation text, strips prefix labels."""
    if not expl_raw:
        return ""
    # Strip व्याख्या: / Explanation: prefix
    text = re.sub(r'^\s*(?:व्याख्या|Explanation)\s*[:–-]\s*', '', expl_raw, flags=re.IGNORECASE)
    return clean_text(text)


# ---------------------------------------------------------------------------
# STEP 8 — Parse question number and body text from <question> raw content
# ---------------------------------------------------------------------------

def parse_question_text(question_raw):
    """
    Returns (q_no_str, body_text).
    Strips the leading number from the body.
    """
    if not question_raw:
        return "", ""

    text = question_raw.strip()

    # Match: 1. / प्रश्न 1 / Q1. / Q1:
    m = re.match(r'^\s*(?:प्रश्न\s*|Q\.?\s*)?(\d+)[.):\s]\s*', text, re.IGNORECASE)
    if m:
        q_no   = m.group(1)
        body   = text[m.end():].strip()
        body   = re.sub(r'^\.+\s*', '', body)
        return q_no, body

    return "", text


# ---------------------------------------------------------------------------
# STEP 9 — Build a complete question dict from raw tag blocks
# ---------------------------------------------------------------------------

def _make_question_dict(raw_q):
    """Convert a raw_q dict (from _group_blocks_into_questions) into the
    full question dict that fill_cell() and the DOCX renderer expect."""

    q_no, q_body = parse_question_text(raw_q["question_raw"])
    if not q_no:
        q_no = raw_q.get("no", "?")

    # Suchi
    suchi_h1 = "सूची-I"
    suchi_h2 = "सूची-II"
    suchi_rows = []
    if raw_q["suchi_raw"]:
        suchi_h1, suchi_h2, suchi_rows = parse_suchi_tagged(raw_q["suchi_raw"])

    # Kathan lines
    kathan_lines = parse_kathan_lines(raw_q["kathan_raw"])

    # Options
    options = parse_options_tagged(raw_q["option_raw"])

    # Answer
    correct = parse_answer_tagged(raw_q["answer_raw"])

    # Explanation
    explanation = parse_expl_tagged(raw_q["expl_raw"])

    # Build final question body:
    # If there were kathan lines, keep question body separate; kathan goes in kathan_lines field.
    # The question body is shown as the main question text.
    final_question = q_body

    return {
        "no":                   q_no,
        "orig_no":              q_no,
        "question":             final_question,
        "kathan_lines":         kathan_lines,   # NEW field for kathan/statement rendering
        "suchi_rows":           suchi_rows,
        "suchi_col_headers":    (suchi_h1, suchi_h2),
        "match_following_rows": [],
        "koot_grid":            {"is_grid": False},
        "options":              options,
        "correct":              correct,
        "explanation":          explanation,
        "explanation_images":   [],
        "metadata":             "",
        "_layout":              None,
    }


# ---------------------------------------------------------------------------
# STEP 10 — Top-level tag-based parse entry point
# ---------------------------------------------------------------------------

def _has_tags(text):
    """Return True if the document uses the tagged format."""
    return bool(_TAG_RE.search(text))


def parse_questions_tagged(doc_or_text, is_text=False):
    """
    Primary parser: tag-based.
    Works for both Hindi and English tagged documents.
    Accepts either a python-docx Document object, or a raw text string
    (pass is_text=True) so the same parser can be reused for .txt uploads.

    NOTE: numbers are kept exactly as written in the source (stored in both
    'no' and 'orig_no'). Final numbering (default / sequenced / customized)
    is applied later, once, after all uploaded files have been merged —
    see apply_numbering_mode().
    """
    full_text = doc_or_text if is_text else _doc_to_text(doc_or_text)
    blocks    = _split_into_tag_blocks(full_text)

    if not blocks:
        return []

    raw_qs    = _group_blocks_into_questions(blocks)
    questions = []
    for idx, raw_q in enumerate(raw_qs):
        q = _make_question_dict(raw_q)
        if not q.get("no"):
            q["no"] = str(idx + 1)
        q["orig_no"] = q["no"]
        questions.append(q)

    return questions


# ---------------------------------------------------------------------------
# FALLBACK — heuristic parsers from original code (kept for non-tagged files)
# ---------------------------------------------------------------------------

def _kw(word):
    return (
        r'(?<![^\u0900-\u097F\s\n।])'
        + word +
        r'(?![^\u0900-\u097F\s\n।\:\-\(\[\u0964\u09650-9१-९])'
    )


def is_standalone_koot(text):
    cleaned = re.sub(r'[\u0900-\u097F]कूट', '', text)
    koot_header_pattern = re.compile(
        r'(?:^|\n)\s*' + _kw('कूट'),
        re.MULTILINE
    )
    return bool(koot_header_pattern.search(cleaned))


def extract_koot_block(text):
    if not is_standalone_koot(text):
        return ""
    cleaned = re.sub(r'[\u0900-\u097F]कूट', '\u0000\u0000\u0000', text)
    pattern = re.compile(
        r'(?:(?:^|\n)\s*)'
        r'(' + _kw('कूट') +
        r'\s*(?:[:\-]\s*|\n|\s+(?=\(|[0-9१-९]))'
        r'.*?)'
        r'(?=\n\s*(?:सही\s*उत्तर|उत्तर\s*[:\-]|व्याख्या\s*[:\-])|$)',
        re.DOTALL | re.MULTILINE
    )
    m = pattern.search(cleaned)
    if not m:
        return ""
    start, end = m.start(1), m.end(1)
    return text[start:end].strip()


def parse_koot_grid(koot_block):
    if not koot_block:
        return {'is_grid': False}
    lines = [l.strip() for l in koot_block.replace('\r', '\n').split('\n') if l.strip()]
    if lines and re.match(r'^कूट', lines[0]):
        lines = lines[1:]
    row_pat = re.compile(
        r'^\s*\(([A-Da-d1-4])\)\s+((?:[a-dA-D\u0966-\u096F]\s*){2,6})\s*$'
    )
    grid_rows   = []
    col_headers = []
    for line in lines:
        if not re.match(r'^\s*\([A-Da-d1-4]\)', line):
            tokens = line.split()
            if all(re.match(r'^[ivxIVX]{1,4}$|^\d$', t) for t in tokens) and len(tokens) >= 2:
                col_headers = tokens
            continue
        m = row_pat.match(line)
        if m:
            key  = f"({m.group(1).upper()})"
            vals = m.group(2).split()
            grid_rows.append((key, vals))
    if len(grid_rows) >= 2:
        return {'is_grid': True, 'col_headers': col_headers, 'rows': grid_rows}
    return {'is_grid': False}


def parse_suchi_table(text):
    """Original heuristic suchi parser — used as fallback only."""
    subtitle1_match = re.search(r'सूची[-–]I\s*\(([^)]*)\)', text)
    subtitle2_match = re.search(r'सूची[-–]II\s*\(([^)]*)\)', text)
    subtitle1 = subtitle1_match.group(1).strip() if subtitle1_match else ""
    subtitle2 = subtitle2_match.group(1).strip() if subtitle2_match else ""
    h1 = f"सूची-I ({subtitle1})" if subtitle1 else "सूची-I"
    h2 = f"सूची-II ({subtitle2})" if subtitle2 else "सूची-II"

    pipe_rows = re.findall(
        r'^([A-D]\.\s+[^|\n]+?)\s*\|\s*(\d+\.\s+.+?)\s*$',
        text, re.MULTILINE
    )
    if pipe_rows:
        rows = [(l.strip(), r.strip()) for l, r in pipe_rows]
        return "", rows, h1, h2

    inline_pat = re.compile(
        r'^\s*(\([^)]+\)\s*[^-–\n]+?)\s*[-–]{1,3}\s*(\([^)]+\)\s*.+?)\s*$',
        re.MULTILINE
    )
    inline_pairs = inline_pat.findall(text)
    if inline_pairs:
        rows = [(left.strip(), right.strip()) for left, right in inline_pairs]
        first_match = inline_pat.search(text)
        header_raw = text[:first_match.start()].strip() if first_match else ""
        header_raw = re.sub(r'सूची[-–]I[I]?\s*(\([^)]*\))?\s*(को)?\s*सूची[-–]II\s*(\([^)]*\))?\s*(से)?\s*सुमेलित\s*कीजिए\s*[:\-]?', '', header_raw).strip()
        header_raw = re.sub(r'सूची[-–]I+\s*(\([^)]*\))?\s*', '', header_raw).strip()
        return header_raw, rows, h1, h2

    return "", [], h1, h2


def parse_match_following_pairs(text, separator_pattern=None):
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    pairs = []
    for line in lines:
        if separator_pattern:
            m = re.match(separator_pattern, line)
            if m:
                pairs.append((m.group(1).strip(), m.group(2).strip()))
            continue
        line = re.sub(r'[–—]+', '-', line)
        line = re.sub(r'-{2,}', '-', line)
        pattern = r'^(\([^)]+\)\s*.*?)\s*-\s*(\([^)]+\)\s*.*)$'
        m = re.match(pattern, line)
        if m:
            pairs.append((m.group(1).strip(), m.group(2).strip()))
    return pairs


# Heuristic Hindi parser (fallback for untagged files)
def parse_questions_heuristic(doc):
    """Original heuristic parser — used only when no tags are found."""
    import io as _io
    questions = []
    current_block = []
    inside_question = False

    def is_question_start(text):
        if not text:
            return False
        text = text.strip()
        return bool(
            re.match(r'^प्रश्न\s+\d+', text) or
            re.match(r'^\d+\.\s+', text)
        )

    def extract_images_from_para(para):
        images = []
        NS = {
            'a':  'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r':  'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        }
        for run in para.runs:
            for blip in run._element.findall('.//a:blip', namespaces=NS):
                r_embed_key = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                rId = blip.get(r_embed_key)
                if not rId:
                    continue
                try:
                    image_part = doc.part.related_parts[rId]
                except KeyError:
                    continue
                img_bytes   = image_part.blob
                width_in = height_in = 1.0
                extent = run._element.find('.//wp:extent', namespaces=NS)
                if extent is not None:
                    width_in  = int(extent.get('cx', 914400)) / 914400.0
                    height_in = int(extent.get('cy', 914400)) / 914400.0
                else:
                    try:
                        pil_img   = PILImage.open(_io.BytesIO(img_bytes))
                        width_in  = pil_img.width  / 96.0
                        height_in = pil_img.height / 96.0
                    except Exception:
                        pass
                images.append((img_bytes, width_in, height_in))
        return images

    for para in doc.paragraphs:
        text   = para.text.strip()
        images = extract_images_from_para(para)

        if re.match(r'^\s*(अथवा|तथा)\s*$', text):
            if current_block:
                q = _process_heuristic_block(current_block)
                if q:
                    q['orig_no'] = q.get('no') or str(len(questions) + 1)
                    q['no']      = q['orig_no']
                    questions.append(q)
            current_block   = []
            inside_question = False
            continue

        if is_question_start(text):
            if current_block:
                q = _process_heuristic_block(current_block)
                if q:
                    q['orig_no'] = q.get('no') or str(len(questions) + 1)
                    q['no']      = q['orig_no']
                    questions.append(q)
            current_block   = [(text, images)]
            inside_question = True
            continue

        if inside_question:
            current_block.append((text, images))

    if current_block:
        q = _process_heuristic_block(current_block)
        if q:
            q['orig_no'] = q.get('no') or str(len(questions) + 1)
            q['no']      = q['orig_no']
            questions.append(q)

    return questions


def _process_heuristic_block(block):
    """Original process_question_block logic."""
    full_text = "\n".join(txt for txt, _ in block).strip()

    q_no = None
    for pattern in [r'प्रश्न\s+(\d+)', r'^(\d+)\.', r'^(\d+)\s+']:
        m = re.search(pattern, full_text)
        if m:
            q_no      = m.group(1)
            full_text = full_text[m.end():].strip()
            full_text = re.sub(r'^\.+\s*', '', full_text)
            break

    if not q_no:
        return None

    ans_match = re.search(
        r'(?:सही उत्तर|उत्तर)\s*[:\-]\s*\(([a-dA-D])\)',
        full_text
    )
    correct = f"({ans_match.group(1).lower()})" if ans_match else ""

    explanation  = ""
    expl_match = re.search(
        r'व्याख्या\s*:\s*(.*?)(?=\n\s*(\d+\.|प्रश्न\s+\d+)|$)',
        full_text, re.DOTALL
    )
    if expl_match:
        explanation = clean_text(expl_match.group(1))

    content = full_text
    if ans_match:
        content = content[:ans_match.start()]
    if expl_match:
        content = content[:expl_match.start()]
    content = content.strip()

    suchi_rows   = []
    suchi_h1     = "सूची-I"
    suchi_h2     = "सूची-II"
    suchi_match = re.search(r'(सूची.*?)(?=(?<!\S)कूट(?!\S)|$)', content, re.DOTALL)
    if suchi_match:
        suchi_block = suchi_match.group(1)
        _, suchi_rows, suchi_h1, suchi_h2 = parse_suchi_table(suchi_block)
        content = content.replace(suchi_match.group(1), "")

    match_following_rows = []

    koot_block = extract_koot_block(content)
    if koot_block:
        content = content.replace(koot_block, "").strip()

    first_opt = re.search(r'\([a-dA-D]\)', content)
    if first_opt:
        question_text = content[:first_opt.start()].strip()
        opts_raw      = content[first_opt.start():]
    else:
        question_text = content
        opts_raw      = ""

    question_text = clean_text(question_text)

    options = []
    if opts_raw:
        matches = re.findall(
            r'\(([a-dA-D])\)\s*(.*?)(?=\([a-dA-D]\)|$)',
            opts_raw, re.DOTALL
        )
        for key, text in matches:
            text = clean_text(text)
            if text:
                options.append({"key": f"({key.lower()})", "text": text.strip()})

    options = options[:4]
    koot_grid = parse_koot_grid(koot_block) if koot_block else {'is_grid': False}

    return {
        "no":                   q_no,
        "question":             question_text,
        "kathan_lines":         [],
        "suchi_rows":           suchi_rows,
        "suchi_col_headers":    (suchi_h1, suchi_h2),
        "match_following_rows": match_following_rows,
        "koot_grid":            koot_grid,
        "options":              options,
        "correct":              correct,
        "explanation":          explanation,
        "explanation_images":   [],
        "metadata":             "",
        "_layout":              None,
    }


# English heuristic fallback (unchanged from original)
_EN_Q_LABEL_RE  = re.compile(r'^\s*Question[-\s]*(\d+)\s*[.:]?\s*$', re.I)
_EN_Q_INLINE_RE = re.compile(r'^\s*Question\s+(\d+)[.):\s]\s*(.+)$', re.I)
_EN_Q_BARE_RE   = re.compile(r'^\s*(\d+)[.)]\s+(.+)$')
_EN_OPT_RE_ENG  = re.compile(r'^\s*\(([a-dA-D])\)\s*(.+)$')
_EN_ANS_RE      = re.compile(r'^\s*(?:Answer|Correct\s*Answer)\s*[-:]\s*\(([a-dA-D])\)', re.I)
_EN_EXPL_HDR_RE = re.compile(r'^\s*(?:Explanation|Explain)\s*[-:]?\s*(.*)', re.I)
_EN_NOTE_RE     = re.compile(r'^\s*\(Note[:\s]', re.I)


def detect_english_paper(doc):
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:30]
    if any(any('\u0900' <= c <= '\u097F' for c in t) for t in texts):
        return False
    has_en_opt = sum(1 for t in texts if _EN_OPT_RE_ENG.match(t))
    has_en_ans = sum(1 for t in texts if _EN_ANS_RE.match(t))
    return has_en_opt >= 2 or has_en_ans >= 1


def _is_en_q_start(text):
    return bool(
        _EN_Q_LABEL_RE.match(text) or
        _EN_Q_INLINE_RE.match(text) or
        _EN_Q_BARE_RE.match(text)
    )


def _collect_en_opts_ans_expl(paras, i, n):
    options    = []
    correct    = ""
    expl_parts = []
    in_expl    = False

    while i < n:
        text = paras[i]
        if text and _is_en_q_start(text):
            break
        if not text:
            i += 1
            if in_expl and expl_parts:
                j = i
                while j < n and not paras[j]:
                    j += 1
                if j >= n or _is_en_q_start(paras[j]):
                    break
            continue
        if _EN_NOTE_RE.match(text):
            i += 1
            continue
        m_opt = _EN_OPT_RE_ENG.match(text)
        if m_opt and not in_expl:
            options.append({"key": f"({m_opt.group(1).lower()})", "text": m_opt.group(2).strip()})
            i += 1
            continue
        m_ans = _EN_ANS_RE.match(text)
        if m_ans and not in_expl:
            correct = f"({m_ans.group(1).lower()})"
            i += 1
            continue
        m_expl = _EN_EXPL_HDR_RE.match(text)
        if m_expl and not in_expl:
            in_expl = True
            tail = m_expl.group(1).strip()
            if tail:
                expl_parts.append(tail)
            i += 1
            continue
        if in_expl:
            expl_parts.append(text)
            i += 1
            continue
        i += 1

    return options, correct, " ".join(expl_parts), i


def parse_english_questions_heuristic(doc):
    paras     = [p.text.strip() for p in doc.paragraphs]
    questions = []
    i         = 0
    n         = len(paras)

    while i < n:
        text = paras[i]
        if not text:
            i += 1
            continue
        m_label = _EN_Q_LABEL_RE.match(text)
        if m_label:
            q_no = m_label.group(1)
            i   += 1
            q_lines = []
            while i < n and paras[i] and not _EN_OPT_RE_ENG.match(paras[i]):
                q_lines.append(paras[i].strip())
                i += 1
            question_text = " ".join(q_lines)
            opts, ans, expl, i = _collect_en_opts_ans_expl(paras, i, n)
            questions.append(_make_en_q(q_no, question_text, opts, ans, expl))
            continue
        m_inline = _EN_Q_INLINE_RE.match(text)
        if m_inline:
            q_no          = m_inline.group(1)
            question_text = m_inline.group(2).strip()
            i            += 1
            opts, ans, expl, i = _collect_en_opts_ans_expl(paras, i, n)
            questions.append(_make_en_q(q_no, question_text, opts, ans, expl))
            continue
        m_bare = _EN_Q_BARE_RE.match(text)
        if m_bare:
            q_no          = m_bare.group(1)
            question_text = m_bare.group(2).strip()
            i            += 1
            opts, ans, expl, i = _collect_en_opts_ans_expl(paras, i, n)
            questions.append(_make_en_q(q_no, question_text, opts, ans, expl))
            continue
        i += 1

    for idx, q in enumerate(questions):
        if not q.get('no'):
            q['no'] = str(idx + 1)
        q['orig_no'] = q['no']
    return questions


def _make_en_q(q_no, question_text, options, correct, explanation):
    return {
        "no":                   q_no,
        "orig_no":              q_no,
        "question":             question_text,
        "kathan_lines":         [],
        "options":              options[:4],
        "correct":              correct,
        "explanation":          explanation,
        "suchi_rows":           [],
        "match_following_rows": [],
        "suchi_col_headers":    ("List-I", "List-II"),
        "koot_grid":            {"is_grid": False},
        "explanation_images":   [],
        "metadata":             "",
        "_layout":              None,
    }


# ---------------------------------------------------------------------------
# MASTER PARSE DISPATCHER — decides tagged vs heuristic
# ---------------------------------------------------------------------------
def parse_questions(doc):
    """
    Master dispatcher:
    1. If document has XML tags → use tag-based parser (handles all 96 suchi variants + kathan)
    2. If English heuristic signals detected → use English heuristic fallback
    3. Otherwise → use Hindi heuristic fallback
    """
    full_text = _doc_to_text(doc)

    if _has_tags(full_text):
        st.info("🏷️ Tagged format detected — using precise tag-based parser")
        return parse_questions_tagged(doc), False   # (questions, is_english)

    _is_english = detect_english_paper(doc)
    if _is_english:
        return parse_english_questions_heuristic(doc), True
    else:
        return parse_questions_heuristic(doc), False


# =============================================================================
# OPTION WIDTH & LAYOUT  (unchanged from original with fixes)
# =============================================================================
def _estimate_option_width_in(opt_text, font_pt):
    text  = opt_text.strip()
    dev   = sum(1 for c in text if '\u0900' <= c <= '\u097F')
    upper = sum(1 for c in text if c.isupper() and c.isascii())
    other = len(text) - dev - upper
    pts   = dev * font_pt * 0.52 + upper * font_pt * 0.65 + other * font_pt * 0.58
    return (pts / 72.0) * 1.15


def layout_options(opts, max_per_line=2, char_limit=68):
    if not opts:
        return []

    content_w      = page_width - left_margin - right_margin
    col_gap        = 0.08 if num_columns == 3 else 0.12
    col_w_in       = (content_w - col_gap * (num_columns - 1)) / num_columns
    cell_margin_in = (40 * 2) / 1440.0
    ans_col_in     = 0.38
    avail_in       = col_w_in - level2_indent - ans_col_in - cell_margin_in * 2 - 0.05

    def fits(opt, n_cols):
        slot = avail_in / n_cols - cell_margin_in
        return _estimate_option_width_in(
            f"{opt['key']} {opt['text']}", opt_font
        ) <= slot

    n = len(opts)

    if n == 4:
        short_opts = all(len(o['text'].strip()) <= 18 for o in opts)
        if short_opts:
            return [[opts[0], opts[1]], [opts[2], opts[3]]]

    capped = min(n, max_per_line)
    if capped >= 2 and all(fits(o, capped) for o in opts):
        return [list(opts[i:i + capped]) for i in range(0, n, capped)]

    result = []
    i = 0
    while i < n:
        placed = False
        for k in range(min(max_per_line, n - i), 1, -1):
            group = opts[i:i + k]
            if all(fits(o, k) for o in group):
                result.append(list(group))
                i += k
                placed = True
                break
        if not placed:
            result.append([opts[i]])
            i += 1

    PASS2_WORDS = 3
    PASS2_CHARS = 25

    def short_enough(opt):
        t = opt['text'].strip()
        return len(t.split()) <= PASS2_WORDS and len(t) <= PASS2_CHARS

    out = []
    j   = 0
    while j < len(result):
        row      = result[j]
        next_row = result[j + 1] if j + 1 < len(result) else None
        if (len(row) == 1
                and next_row is not None
                and len(next_row) == 1
                and short_enough(row[0])
                and short_enough(next_row[0])):
            out.append([row[0], next_row[0]])
            j += 2
        else:
            out.append(row)
            j += 1

    return out


# =============================================================================
# DOCX HELPERS  (unchanged)
# =============================================================================
def set_spacing(para, line_pts, after_pts=0, before_pts=0):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    s = OxmlElement('w:spacing')
    s.set(qn('w:line'),     str(int(line_pts * 20)))
    s.set(qn('w:lineRule'), 'atLeast')
    s.set(qn('w:before'),   str(int(before_pts * 20)))
    s.set(qn('w:after'),    str(int(after_pts  * 20)))
    pPr.append(s)


def set_char_spacing(run, spacing_pt):
    if spacing_pt > 0:
        rPr = run._r.get_or_add_rPr()
        sp  = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(int(spacing_pt * 20)))
        rPr.append(sp)


def set_paragraph_background(para, color_rgb):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_rgb)
    pPr = para._p.get_or_add_pPr()
    pPr.append(shd)


def _apply_ind(para, left_twips, first_twips):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:ind')):
        pPr.remove(old)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left_twips))
    if first_twips != 0:
        ind.set(qn('w:firstLine'), str(first_twips))
    pPr.append(ind)


def set_two_level_indent(para, l1_in, l2_in):
    _apply_ind(para, int(l2_in * 1440), int((l1_in - l2_in) * 1440))


def set_left_indent(para, left_in):
    _apply_ind(para, int(left_in * 1440), 0)


def apply_font_to_run(run):
    run.font.name = FONT_DOCX
    rPr    = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    FONT_DOCX)
    rFonts.set(qn('w:hAnsi'),    FONT_DOCX)
    rFonts.set(qn('w:eastAsia'), FONT_DOCX)
    rFonts.set(qn('w:cs'),       FONT_DOCX)


def add_run(para, text, bold=False, size_pt=8, italic=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.font.size = Pt(size_pt)
    apply_font_to_run(r)
    if char_spacing > 0:
        set_char_spacing(r, char_spacing)
    return r


# =============================================================================
# NEW: KATHAN RENDERER — renders kathan_lines as indented block in DOCX
# =============================================================================
def add_kathan_block(container, kathan_lines):
    """
    Renders each kathan/statement line as an indented paragraph.
    Applies kathan_font, kathan_bold, kathan_indent_extra styling from sidebar.
    Optional light grey background if kathan_bg is True.
    """
    if not kathan_lines:
        return
    for line in kathan_lines:
        p = container.add_paragraph()
        p.paragraph_format.left_indent = Inches(level2_indent + kathan_indent_extra)
        if kathan_bg:
            set_paragraph_background(p, "F5F5F5")
        add_run(p, line, bold=kathan_bold, size_pt=kathan_font)
        set_spacing(p, line_pts=line_spacing, after_pts=0)


# =============================================================================
# सूची TABLE  (unchanged)
# =============================================================================
def add_suchi_table(container, suchi_rows, col_width_in, col_headers=("सूची-I", "सूची-II")):
    if not suchi_rows:
        return
    half_dxa  = int((col_width_in - level2_indent) * 1440 / 2)
    total_dxa = half_dxa * 2

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW  = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(total_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'),    str(int(level2_indent * 1440)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    '0')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)
    tbl.append(tblPr)

    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(2):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(half_dxa))
        tblGrid.append(gc)
    tbl.append(tblGrid)

    def make_tc(text_content, width_dxa, bold=False):
        tc   = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        tcMar = OxmlElement('w:tcMar')
        for edge, val in [('top', '0'), ('left', '40'), ('bottom', '0'), ('right', '80')]:
            m = OxmlElement(f'w:{edge}')
            m.set(qn('w:w'),    val)
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        tc.append(tcPr)
        p   = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        sp  = OxmlElement('w:spacing')
        sp.set(qn('w:line'),     str(int(line_spacing * 20)))
        sp.set(qn('w:lineRule'), 'atLeast')
        sp.set(qn('w:before'),   '0')
        sp.set(qn('w:after'),    '0')
        pPr.append(sp)
        p.append(pPr)
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        sz  = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(q_font * 2)))
        rPr.append(sz)
        rF = OxmlElement('w:rFonts')
        rF.set(qn('w:ascii'),   FONT_DOCX)
        rF.set(qn('w:hAnsi'),   FONT_DOCX)
        rF.set(qn('w:cs'),      FONT_DOCX)
        rPr.insert(0, rF)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text_content
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    def make_tc_shaded(text_content, width_dxa, shade="D9D9D9"):
        tc   = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  shade)
        tcPr.append(shd)
        tcMar = OxmlElement('w:tcMar')
        for edge, val in [('top', '0'), ('left', '40'), ('bottom', '0'), ('right', '80')]:
            m = OxmlElement(f'w:{edge}')
            m.set(qn('w:w'),    val)
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        tc.append(tcPr)
        p   = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        sp  = OxmlElement('w:spacing')
        sp.set(qn('w:line'),     str(int(line_spacing * 20)))
        sp.set(qn('w:lineRule'), 'atLeast')
        sp.set(qn('w:before'),   '0')
        sp.set(qn('w:after'),    '0')
        pPr.append(sp)
        p.append(pPr)
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))
        sz  = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(q_font * 2)))
        rPr.append(sz)
        rF = OxmlElement('w:rFonts')
        rF.set(qn('w:ascii'),   FONT_DOCX)
        rF.set(qn('w:hAnsi'),   FONT_DOCX)
        rF.set(qn('w:cs'),      FONT_DOCX)
        rPr.insert(0, rF)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text_content
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    h1_label, h2_label = col_headers
    tr_hdr   = OxmlElement('w:tr')
    trPr_hdr = OxmlElement('w:trPr')
    trH_hdr  = OxmlElement('w:trHeight')
    trH_hdr.set(qn('w:val'),   str(int(line_spacing * 20)))
    trH_hdr.set(qn('w:hRule'), 'atLeast')
    trPr_hdr.append(trH_hdr)
    tr_hdr.append(trPr_hdr)
    tr_hdr.append(make_tc_shaded(h1_label, half_dxa))
    tr_hdr.append(make_tc_shaded(h2_label, half_dxa))
    tbl.append(tr_hdr)

    for left_text, right_text in suchi_rows:
        tr   = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        trH  = OxmlElement('w:trHeight')
        trH.set(qn('w:val'),   str(int(line_spacing * 20)))
        trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)
        tr.append(trPr)
        tr.append(make_tc(left_text,  half_dxa))
        tr.append(make_tc(right_text, half_dxa))
        tbl.append(tr)

    last_para = container.paragraphs[-1]._element if container.paragraphs else None
    if last_para is not None:
        last_para.addnext(tbl)
    else:
        container._element.body.append(tbl)


# =============================================================================
# MATCH THE FOLLOWING TABLE  (unchanged)
# =============================================================================
def add_match_following_table(container, rows, col_width_in):
    if not rows:
        return
    avail_in  = col_width_in - level2_indent
    left_pct  = mtf_left_pct / 100.0
    left_dxa  = int(avail_in * 1440 * left_pct)
    right_dxa = int(avail_in * 1440 * (1 - left_pct))
    total_dxa = left_dxa + right_dxa
    pad_str   = str(int(mtf_cell_pad))
    inside_v_val = 'single' if mtf_inside_v else 'nil'

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableNormal')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(total_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'),    str(int(level2_indent * 1440)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    b_iv = OxmlElement('w:insideV')
    b_iv.set(qn('w:val'),   inside_v_val)
    b_iv.set(qn('w:sz'),    '4')
    b_iv.set(qn('w:space'), '0')
    b_iv.set(qn('w:color'), 'AAAAAA')
    tblBorders.append(b_iv)
    tblPr.append(tblBorders)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    '0')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)
    tblLook = OxmlElement('w:tblLook')
    tblLook.set(qn('w:val'),         '0000')
    tblLook.set(qn('w:firstRow'),    '0')
    tblLook.set(qn('w:lastRow'),     '0')
    tblLook.set(qn('w:firstColumn'), '0')
    tblLook.set(qn('w:lastColumn'),  '0')
    tblLook.set(qn('w:noHBand'),     '1')
    tblLook.set(qn('w:noVBand'),     '1')
    tblPr.append(tblLook)
    tbl.append(tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    for dxa in [left_dxa, right_dxa]:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(dxa))
        tblGrid.append(gc)
    tbl.append(tblGrid)

    def make_tc(text_content, width_dxa, shaded=False):
        tc   = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        if shaded:
            shd_el = OxmlElement('w:shd')
            shd_el.set(qn('w:val'),   'clear')
            shd_el.set(qn('w:color'), 'auto')
            shd_el.set(qn('w:fill'),  'D9D9D9')
            tcPr.append(shd_el)
        tcMar = OxmlElement('w:tcMar')
        for edge, val in [('top','0'),('left', pad_str),('bottom','0'),('right', pad_str)]:
            m = OxmlElement(f'w:{edge}')
            m.set(qn('w:w'),    val)
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        tc.append(tcPr)
        p   = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'),  '0')
        ind.set(qn('w:right'), '0')
        pPr.append(ind)
        sp  = OxmlElement('w:spacing')
        sp.set(qn('w:line'),     str(int(line_spacing * 20)))
        sp.set(qn('w:lineRule'), 'atLeast')
        sp.set(qn('w:before'),   '0')
        sp.set(qn('w:after'),    '0')
        pPr.append(sp)
        p.append(pPr)
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if shaded:
            rPr.append(OxmlElement('w:b'))
        sz  = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(opt_font * 2)))
        rPr.append(sz)
        rF  = OxmlElement('w:rFonts')
        rF.set(qn('w:ascii'),   FONT_DOCX)
        rF.set(qn('w:hAnsi'),   FONT_DOCX)
        rF.set(qn('w:cs'),      FONT_DOCX)
        rPr.insert(0, rF)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text_content
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    def make_row(left_text, right_text, shaded=False):
        tr   = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        trH  = OxmlElement('w:trHeight')
        trH.set(qn('w:val'),   str(int(line_spacing * 20)))
        trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)
        tr.append(trPr)
        tr.append(make_tc(left_text,  left_dxa,  shaded=shaded))
        tr.append(make_tc(right_text, right_dxa, shaded=shaded))
        return tr

    if mtf_show_header:
        tbl.append(make_row(mtf_h1, mtf_h2, shaded=True))
    for left_text, right_text in rows:
        tbl.append(make_row(left_text, right_text))

    last_para = container.paragraphs[-1]._element if container.paragraphs else None
    if last_para is not None:
        last_para.addnext(tbl)
    else:
        container._element.body.append(tbl)


# =============================================================================
# कूट GRID TABLE  (unchanged)
# =============================================================================
def add_koot_grid_table(container, koot_grid, col_width_in, correct_ans):
    col_headers = koot_grid.get('col_headers', [])
    grid_rows   = koot_grid.get('rows', [])
    if not grid_rows:
        return

    n_val_cols  = max(len(r[1]) for r in grid_rows)
    n_total_col = 1 + n_val_cols
    avail_in   = col_width_in - level2_indent
    key_col_in = 0.30
    val_col_in = (avail_in - key_col_in) / n_val_cols
    key_dxa    = int(key_col_in * 1440)
    val_dxa    = int(val_col_in * 1440)
    total_dxa  = key_dxa + val_dxa * n_val_cols

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblW  = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(total_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'),    str(int(level2_indent * 1440)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '40')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)
    tbl.append(tblPr)

    tblGrid = OxmlElement('w:tblGrid')
    gc = OxmlElement('w:gridCol')
    gc.set(qn('w:w'), str(key_dxa))
    tblGrid.append(gc)
    for _ in range(n_val_cols):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(val_dxa))
        tblGrid.append(gc)
    tbl.append(tblGrid)

    def make_cell(text, width_dxa, bold=False, align='center', shade=None):
        tc   = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        if shade:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  shade)
            tcPr.append(shd)
        tc.append(tcPr)
        p   = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc  = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        pPr.append(jc)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:line'),     str(int(line_spacing * 20)))
        sp.set(qn('w:lineRule'), 'atLeast')
        sp.set(qn('w:before'),   '0')
        sp.set(qn('w:after'),    '0')
        pPr.append(sp)
        p.append(pPr)
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(opt_font * 2)))
        rPr.append(sz)
        rF = OxmlElement('w:rFonts')
        rF.set(qn('w:ascii'),   FONT_DOCX)
        rF.set(qn('w:hAnsi'),   FONT_DOCX)
        rF.set(qn('w:cs'),      FONT_DOCX)
        rPr.insert(0, rF)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    def make_row(cells):
        tr   = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        trH  = OxmlElement('w:trHeight')
        trH.set(qn('w:val'),   str(int(line_spacing * 20)))
        trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)
        tr.append(trPr)
        for cell in cells:
            tr.append(cell)
        return tr

    if col_headers:
        cells = [make_cell("", key_dxa, bold=True, shade="D9D9D9")]
        for h in col_headers[:n_val_cols]:
            cells.append(make_cell(h, val_dxa, bold=True, shade="D9D9D9"))
        for _ in range(n_val_cols - len(col_headers)):
            cells.append(make_cell("", val_dxa, shade="D9D9D9"))
        tbl.append(make_row(cells))

    correct_key = (correct_ans or "").strip("() ").upper() if correct_ans else ""
    for key_label, val_list in grid_rows:
        is_correct = key_label.strip("() ").upper() == correct_key
        cells = [make_cell(key_label, key_dxa, bold=is_correct, align='left')]
        for i in range(n_val_cols):
            v = val_list[i] if i < len(val_list) else ""
            cells.append(make_cell(v, val_dxa, bold=is_correct))
        tbl.append(make_row(cells))

    last_para = container.paragraphs[-1]._element if container.paragraphs else None
    if last_para is not None:
        last_para.addnext(tbl)
    else:
        container._element.body.append(tbl)


# =============================================================================
# OPTIONS TABLE  (unchanged from original with FIX 3)
# =============================================================================
def add_options_table(container, option_groups, col_width_in, correct_ans, anchor=None):
    if not option_groups:
        return

    all_opts = [o for row in option_groups for o in row]

    is_koot = any(
        re.search(r'[①②③④⑤⑥⑦⑧]', o['text']) or
        re.search(r'\b[ivxIVX]{1,4}\s*[-–]\s*[ivxIVX]{1,4}', o['text']) or
        re.search(r'\b[ivxIVX]{1,4}\s*[-–]\s*[a-dA-D]', o['text']) or
        re.search(r'\(\s*[ivxIVX]{1,4}\s*\)\s*[-–]', o['text'])
        for o in all_opts
    )

    if is_koot:
        for idx, opt in enumerate(all_opts):
            is_last = (idx == len(all_opts) - 1)
            p = container.add_paragraph()
            p.paragraph_format.left_indent = Inches(level2_indent + 0.1)
            add_run(p, f"{opt['key']} {opt['text']}", bold=opt_bold, size_pt=opt_font)
            set_spacing(p, line_pts=line_spacing, after_pts=0)
            if is_last and show_correct_inline and correct_ans:
                r_ans = p.add_run(f"  {correct_ans}")
                r_ans.bold      = True
                r_ans.font.size = Pt(opt_font)
                apply_font_to_run(r_ans)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return

    avail_in     = col_width_in - level2_indent
    ans_col_in   = 0.38
    opt_total_in = avail_in - ans_col_in
    ans_col_dxa  = int(ans_col_in * 1440)
    ind_dxa      = int(level2_indent * 1440)
    total_rows   = len(option_groups)

    def make_opt_tc(text, width_dxa, bold=False, align='left'):
        tc   = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', '0'), ('left', '40'), ('bottom', '0'), ('right', '40')]:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'),    val)
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        tc.append(tcPr)
        p   = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc  = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        pPr.append(jc)
        sp  = OxmlElement('w:spacing')
        sp.set(qn('w:line'),     str(int(line_spacing * 20)))
        sp.set(qn('w:lineRule'), 'atLeast')
        sp.set(qn('w:before'),   '0')
        sp.set(qn('w:after'),    '0')
        pPr.append(sp)
        p.append(pPr)
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold or opt_bold:
            rPr.append(OxmlElement('w:b'))
        sz  = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(opt_font * 2)))
        rPr.append(sz)
        rF  = OxmlElement('w:rFonts')
        rF.set(qn('w:ascii'),   FONT_DOCX)
        rF.set(qn('w:hAnsi'),   FONT_DOCX)
        rF.set(qn('w:cs'),      FONT_DOCX)
        rPr.insert(0, rF)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    def make_row_table(group, is_last_group):
        n_cols = len(group)
        if n_cols == 1 and not is_last_group:
            col_w_dxa   = int(avail_in * 1440)
            total_dxa   = col_w_dxa
            include_ans = False
        elif n_cols == 1 and is_last_group:
            col_w_dxa   = int(opt_total_in * 1440 * 0.96)
            total_dxa   = col_w_dxa + ans_col_dxa
            include_ans = True
        else:
            col_w_dxa   = int(opt_total_in / n_cols * 1440 * 0.96)
            total_dxa   = col_w_dxa * n_cols + ans_col_dxa
            include_ans = True

        tbl = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        tblW  = OxmlElement('w:tblW')
        tblW.set(qn('w:w'),    str(total_dxa))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'),    str(ind_dxa))
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tblBorders.append(b)
        tblPr.append(tblBorders)
        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ['top', 'left', 'bottom', 'right']:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'),    '0')
            m.set(qn('w:type'), 'dxa')
            tblCellMar.append(m)
        tblPr.append(tblCellMar)
        tbl.append(tblPr)

        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(n_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(col_w_dxa))
            tblGrid.append(gc)
        if include_ans:
            gc_ans = OxmlElement('w:gridCol')
            gc_ans.set(qn('w:w'), str(ans_col_dxa))
            tblGrid.append(gc_ans)
        tbl.append(tblGrid)

        tr   = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        trH  = OxmlElement('w:trHeight')
        trH.set(qn('w:val'),   str(int(line_spacing * 20)))
        trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)
        tr.append(trPr)

        for opt in group:
            tr.append(make_opt_tc(f"{opt['key']} {opt['text']}", col_w_dxa))

        if include_ans:
            if is_last_group and show_correct_inline and correct_ans:
                tr.append(make_opt_tc(correct_ans, ans_col_dxa, bold=True, align='right'))
            else:
                tr.append(make_opt_tc("", ans_col_dxa))

        tbl.append(tr)
        return tbl

    # ── Insert option tables in correct order using advancing anchor ─────────
    current_anchor = anchor
    for row_idx, group in enumerate(option_groups):
        is_last = (row_idx == total_rows - 1)
        tbl = make_row_table(group, is_last)
        if current_anchor is not None:
            current_anchor.addnext(tbl)
            current_anchor = tbl
        else:
            container._element.body.append(tbl)
            current_anchor = tbl

# -=-------  Fill cell  

def fill_cell(container, q, include_metadata=False):
    content_width = page_width - left_margin - right_margin
    col_gap   = 0.08 if num_columns == 3 else 0.12
    col_width = (content_width - col_gap * (num_columns - 1)) / num_columns

    if q.get('is_separator'):
        p_sep = container.add_paragraph()
        set_paragraph_background(p_sep, "3B3B3B")
        p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_sep, f"📄  {q.get('text', '')}", bold=True, size_pt=max(q_font, 10.0))
        set_spacing(p_sep, line_pts=line_spacing, before_pts=4, after_pts=4)
        return

    # ── Question paragraph ──────────────────────────────────────────────────
    p_q = container.add_paragraph()
    p_q.paragraph_format.left_indent       = Inches(level2_indent)
    p_q.paragraph_format.first_line_indent = Inches(level1_indent - level2_indent)

    tab_stops = p_q.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(level2_indent),   WD_TAB_ALIGNMENT.LEFT)
    tab_stops.add_tab_stop(Inches(col_width - 0.2), WD_TAB_ALIGNMENT.LEFT)

    display_question = q['question']
    add_run(p_q, f"{q['no']}. ", bold=True,   size_pt=q_font)
    add_run(p_q, display_question, bold=q_bold, size_pt=q_font)
    set_spacing(p_q, line_pts=line_spacing, after_pts=para_spacing)

    if include_metadata and q.get('metadata'):
        p_meta = container.add_paragraph()
        p_meta.paragraph_format.left_indent = Inches(level2_indent)
        p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_meta = p_meta.add_run(q['metadata'])
        r_meta.italic    = True
        r_meta.font.size = Pt(max(q_font - 1.0, 5.0))
        apply_font_to_run(r_meta)
        set_spacing(p_meta, line_pts=line_spacing, after_pts=0)

    # ── Kathan / Statement lines ────────────────────────────────────────────
    if q.get('kathan_lines'):
        add_kathan_block(container, q['kathan_lines'])

    # ── Suchi table ─────────────────────────────────────────────────────────
    if q.get('suchi_rows'):
        add_suchi_table(container, q['suchi_rows'], col_width,
                        col_headers=q.get('suchi_col_headers', ("सूची-I", "सूची-II")))

    # ── Match the following table ────────────────────────────────────────────
    if q.get('match_following_rows'):
        add_match_following_table(container, q['match_following_rows'], col_width)

    # ── Fresh anchor paragraph placed AFTER all tables ───────────────────────
    # We add a real zero-height paragraph so addnext() in add_options_table
    # has a guaranteed position that is AFTER kathan/suchi/match XML tables.
    p_anchor = container.add_paragraph()
    set_spacing(p_anchor, line_pts=0, after_pts=0)
    anchor_el = p_anchor._element

    # ── Koot grid table OR options table ─────────────────────────────────────
    kg = q.get('koot_grid', {})
    if kg.get('is_grid'):
        anchor_el.getparent().remove(anchor_el)
        add_koot_grid_table(container, kg, col_width, q['correct'])
    else:
        option_groups = q.get('_layout', layout_options(
            q['options'], max_per_line=opts_per_line, char_limit=opt_char_limit
        ))
        add_options_table(container, option_groups, col_width, q['correct'],
                          anchor=anchor_el)
        anchor_el.getparent().remove(anchor_el)

    # ── Explanation ──────────────────────────────────────────────────────────
    if q['explanation']:
        p_expl = container.add_paragraph()
        p_expl.paragraph_format.left_indent       = Inches(level2_indent)
        p_expl.paragraph_format.first_line_indent = Inches(level1_indent - level2_indent)
        if expl_bg:
            set_paragraph_background(p_expl, "E6E6E6")
        _label  = expl_prefix if expl_prefix.strip() else "व्याख्या"
        _prefix = f"➤ {_label}: " if expl_bullet else f"{_label}: "
        add_run(p_expl, _prefix,          bold=True, size_pt=expl_font)
        add_run(p_expl, q['explanation'],             size_pt=expl_font)
        set_spacing(p_expl, line_pts=line_spacing, before_pts=0, after_pts=para_spacing * 2)
        pPr_e = p_expl._p.get_or_add_pPr()
        ctxSp = OxmlElement('w:contextualSpacing')
        ctxSp.set(qn('w:val'), '1')
        pPr_e.append(ctxSp)

    # ── Explanation images ────────────────────────────────────────────────────
    for idx, (img_bytes, width_in, height_in) in enumerate(q.get('explanation_images', [])):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                tmp.write(img_bytes)
                tmp_path = tmp.name
            max_img_w = col_width - level2_indent - 0.2
            img_w     = min(width_in if width_in > 0 else 1.5, max_img_w)
            p_img     = container.add_paragraph()
            p_img.paragraph_format.left_indent = Inches(level2_indent)
            p_img.add_run().add_picture(tmp_path, width=Inches(img_w))
            os.unlink(tmp_path)
            set_spacing(p_img, line_pts=line_spacing, before_pts=0, after_pts=para_spacing)
        except Exception:
            p_ph = container.add_paragraph()
            p_ph.paragraph_format.left_indent = Inches(level2_indent)
            add_run(p_ph, "[चित्र यहाँ संलग्न करें]", italic=True, size_pt=expl_font)
            if expl_bg:
                set_paragraph_background(p_ph, "E6E6E6")
            set_spacing(p_ph, line_pts=line_spacing, after_pts=para_spacing)

    if show_separator:
        p_sep = container.add_paragraph()
        set_spacing(p_sep, line_pts=line_spacing, after_pts=2)




def estimate_q_lines(q):
    if q.get('is_separator'):
        return 2
    lines  = 1
    lines += len(q.get('kathan_lines', []))
    lines += len(layout_options(q['options'], max_per_line=opts_per_line, char_limit=opt_char_limit))
    if q['explanation']:
        lines += 1
    lines += len(q.get('suchi_rows', [])) + (1 if q.get('suchi_rows') else 0)
    lines += len(q.get('explanation_images', [])) * 3
    return lines


# =============================================================================
# TWO-PASS GLOBAL LAYOUT PRECOMPUTATION  (unchanged)
# =============================================================================
def precompute_layouts(questions):
    for q in questions:
        q['_layout'] = layout_options(
            q.get('options', []),
            max_per_line=opts_per_line,
            char_limit=opt_char_limit
        )

    WORD_LIMIT = 3
    CHAR_LIMIT = 25

    def can_pair(opt):
        text = opt['text'].strip()
        return len(text.split()) <= WORD_LIMIT and len(text) <= CHAR_LIMIT

    for q in questions:
        rows = q['_layout']
        if len(rows) == 1:
            continue
        out = []
        i   = 0
        while i < len(rows):
            row = rows[i]
            if (len(row) == 1
                    and i + 1 < len(rows)
                    and len(rows[i + 1]) == 1
                    and can_pair(row[0])
                    and can_pair(rows[i + 1][0])):
                out.append([row[0], rows[i + 1][0]])
                i += 2
            else:
                out.append(row)
                i += 1
        q['_layout'] = out


# =============================================================================
# PAGE GENERATION  (unchanged)
# =============================================================================
def generate_multi_page_docx(questions, chapter_title):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Inches(page_width)
    sec.page_height   = Inches(page_height)
    sec.top_margin    = Inches(top_margin)
    sec.bottom_margin = Inches(bottom_margin)
    sec.left_margin   = Inches(left_margin)
    sec.right_margin  = Inches(right_margin)

    sectPr = sec._sectPr
    W_NS   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    cols_list = sectPr.findall(f'{{{W_NS}}}cols')
    if cols_list:
        cols = cols_list[0]
    else:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'),   str(num_columns))
    cols.set(qn('w:space'), "300")

    BG            = "E6E6E6"
    LOGO_H_INCHES = 0.22
    LOGO_W_INCHES = LOGO_H_INCHES * (922 / 376)
    LOGO_PATH     = "logo.png"

    HEADER_HEIGHT_IN = LOGO_H_INCHES + (header_top_padding / 1440.0)
    GAP_5MM_IN       = 0.197
    sec.header_distance = Inches(max(0.05, top_margin - HEADER_HEIGHT_IN - GAP_5MM_IN))

    titlePg = OxmlElement('w:titlePg')
    sectPr.append(titlePg)

    page_w_dxa = int(page_width  * 1440)
    lm_dxa     = int(left_margin * 1440)
    rm_dxa     = int(right_margin * 1440)
    total_dxa  = page_w_dxa - lm_dxa - rm_dxa

    PAGE_COL_DXA = int(0.38 * 1440)
    LOGO_COL_DXA = int(LOGO_W_INCHES * 1440)
    MID_COL_DXA  = total_dxa - PAGE_COL_DXA - LOGO_COL_DXA

    clean_chapter = chapter_heading.strip() if chapter_heading.strip() else re.sub(r'\*+', '', chapter_title).strip()

    def make_header_cell(width_dxa, top_pad_twips=None):
        pad = top_pad_twips if top_pad_twips is not None else header_top_padding
        tc   = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_dxa))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  BG)
        tcPr.append(shd)
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        tcMar = OxmlElement('w:tcMar')
        for edge, val in [
            ('top',    str(pad)),
            ('left',   '60'),
            ('bottom', '60'),
            ('right',  '60'),
        ]:
            m = OxmlElement(f'w:{edge}')
            m.set(qn('w:w'),    val)
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        tc.append(tcPr)
        return tc

    def make_text_para(text, align='center', bold=True, font_pt=None):
        """Return a bare <w:p> XML element with styled text."""
        fpt = font_pt or header_font
        p   = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc  = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        pPr.append(jc)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'),  '0')
        pPr.append(sp)
        p.append(pPr)
        r   = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold and header_bold:
            rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(fpt * 2)))
        rPr.append(sz)
        rF = OxmlElement('w:rFonts')
        rF.set(qn('w:ascii'),   FONT_DOCX)
        rF.set(qn('w:hAnsi'),   FONT_DOCX)
        rF.set(qn('w:cs'),      FONT_DOCX)
        rPr.insert(0, rF)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def make_page_num_para(align='left', font_pt=None):
        """Return a <w:p> with an auto PAGE field."""
        fpt = font_pt or header_font
        p   = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc  = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        pPr.append(jc)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'),  '0')
        pPr.append(sp)
        p.append(pPr)
        for ftype, itext in [('begin', None), ('instr', ' PAGE '), ('end', None)]:
            r   = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rPr.append(OxmlElement('w:b'))
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(fpt * 2)))
            rPr.append(sz)
            r.append(rPr)
            if ftype == 'instr':
                el = OxmlElement('w:instrText')
                el.set(qn('xml:space'), 'preserve')
                el.text = itext
                r.append(el)
            else:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), ftype)
                r.append(fc)
            p.append(r)
        return p

    def make_logo_tc(logo_path, logo_w_in, logo_h_in, col_dxa, header_part):
        tc   = make_header_cell(col_dxa)
        p_el = OxmlElement('w:p')
        pPr  = OxmlElement('w:pPr')
        jc   = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'),  '0')
        pPr.append(sp)
        p_el.append(pPr)
        tc.append(p_el)

        if os.path.exists(logo_path):
            from docx.opc.part    import Part
            from docx.opc.packuri import PackURI
            from lxml import etree

            with open(logo_path, 'rb') as f:
                img_data = f.read()

            ext_map  = {'.png': 'image/png', '.jpg': 'image/jpeg',
                        '.jpeg': 'image/jpeg', '.gif': 'image/gif', '.bmp': 'image/bmp'}
            ext      = os.path.splitext(logo_path)[1].lower()
            ct       = ext_map.get(ext, 'image/png')
            partname = PackURI(f'/word/media/logo_hdr{ext}')
            img_part = Part(partname, ct, img_data)
            REL_TYPE = ('http://schemas.openxmlformats.org/officeDocument/'
                        '2006/relationships/image')
            rId = header_part.relate_to(img_part, REL_TYPE)

            cx = int(logo_w_in * 914400)
            cy = int(logo_h_in * 914400)

            WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            PIC = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
            R   = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

            r_el   = etree.SubElement(p_el, f'{{{W}}}r')
            draw   = etree.SubElement(r_el, f'{{{W}}}drawing')
            inline = etree.SubElement(draw, f'{{{WP}}}inline',
                                      distT='0', distB='0', distL='0', distR='0')
            etree.SubElement(inline, f'{{{WP}}}extent',       cx=str(cx), cy=str(cy))
            etree.SubElement(inline, f'{{{WP}}}effectExtent', l='0', t='0', r='0', b='0')
            etree.SubElement(inline, f'{{{WP}}}docPr',        id='200', name='logo_hdr')
            cNv = etree.SubElement(inline, f'{{{WP}}}cNvGraphicFramePr')
            etree.SubElement(cNv,    f'{{{A}}}graphicFrameLocks', noChangeAspect='1')
            graphic     = etree.SubElement(inline,     f'{{{A}}}graphic')
            graphicData = etree.SubElement(graphic,    f'{{{A}}}graphicData', uri=PIC)
            pic_el      = etree.SubElement(graphicData, f'{{{PIC}}}pic')
            nvPicPr     = etree.SubElement(pic_el,     f'{{{PIC}}}nvPicPr')
            etree.SubElement(nvPicPr, f'{{{PIC}}}cNvPr', id='0', name='logo')
            etree.SubElement(nvPicPr, f'{{{PIC}}}cNvPicPr')
            blipFill = etree.SubElement(pic_el,      f'{{{PIC}}}blipFill')
            blip     = etree.SubElement(blipFill,    f'{{{A}}}blip')
            blip.set(f'{{{R}}}embed', rId)
            stretch  = etree.SubElement(blipFill,    f'{{{A}}}stretch')
            etree.SubElement(stretch, f'{{{A}}}fillRect')
            spPr     = etree.SubElement(pic_el,      f'{{{PIC}}}spPr')
            xfrm     = etree.SubElement(spPr,        f'{{{A}}}xfrm')
            etree.SubElement(xfrm,  f'{{{A}}}off',   x='0', y='0')
            etree.SubElement(xfrm,  f'{{{A}}}ext',   cx=str(cx), cy=str(cy))
            prstGeom = etree.SubElement(spPr,        f'{{{A}}}prstGeom', prst='rect')
            etree.SubElement(prstGeom, f'{{{A}}}avLst')
        else:
            r   = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rPr.append(OxmlElement('w:b'))
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(header_font * 2)))
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = 'RBD'
            r.append(t)
            p_el.append(r)

        return tc

    def text_cell(width_dxa, para_el):
        tc = make_header_cell(width_dxa)
        tc.append(para_el)
        return tc

    def build_header_row(col_widths, cell_elements):
        tbl   = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        tblW  = OxmlElement('w:tblW')
        tblW.set(qn('w:w'),    str(sum(col_widths)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tblBorders.append(b)
        tblPr.append(tblBorders)
        tbl.append(tblPr)
        tblGrid = OxmlElement('w:tblGrid')
        for dxa in col_widths:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(dxa))
            tblGrid.append(gc)
        tbl.append(tblGrid)
        tr   = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        trH  = OxmlElement('w:trHeight')
        min_row_h = int(LOGO_H_INCHES * 1440)
        trH.set(qn('w:val'),   str(min_row_h))
        trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)
        tr.append(trPr)
        for tc_el in cell_elements:
            tr.append(tc_el)
        tbl.append(tr)
        return tbl

    P1_LEFT_DXA = PAGE_COL_DXA
    P1_MID_DXA  = total_dxa - PAGE_COL_DXA - LOGO_COL_DXA

    first_header = sec.first_page_header
    for p in list(first_header.paragraphs):
        p._element.getparent().remove(p._element)

    tbl_first = build_header_row(
        [P1_LEFT_DXA, P1_MID_DXA, LOGO_COL_DXA],
        [
            text_cell(P1_LEFT_DXA,  make_text_para("", align='left')),
            text_cell(P1_MID_DXA,   make_text_para(clean_chapter, align='center', font_pt=header_font + 1)),
            make_logo_tc(LOGO_PATH, LOGO_W_INCHES, LOGO_H_INCHES, LOGO_COL_DXA, sec.first_page_header.part),
        ]
    )
    first_header._element.insert(0, tbl_first)

    header = sec.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    tbl_main = build_header_row(
        [PAGE_COL_DXA, MID_COL_DXA, LOGO_COL_DXA],
        [
            text_cell(PAGE_COL_DXA, make_page_num_para(align='left')),
            text_cell(MID_COL_DXA,  make_text_para(clean_chapter, align='center')),
            make_logo_tc(LOGO_PATH, LOGO_W_INCHES, LOGO_H_INCHES, LOGO_COL_DXA, sec.header.part),
        ]
    )
    header._element.insert(0, tbl_main)

    for q in questions:
        fill_cell(doc, q, include_metadata=include_metadata)

    return doc


# =============================================================================
# HTML PREVIEW — updated to include kathan_lines rendering
# =============================================================================
def render_q_preview(q):
    if q.get('is_separator'):
        return f"""<div class="qblock" style="break-inside:avoid;page-break-inside:avoid;">
  <div style="background:#3B3B3B;color:#fff;font-weight:bold;text-align:center;
    padding:6px;margin:8px 0;border-radius:3px;">📄 {q.get('text', '')}</div>
</div>"""
    l1px = level1_indent * 96
    l2px = level2_indent * 96
    kl_px = (level2_indent + kathan_indent_extra) * 96

    option_groups = layout_options(q['options'], max_per_line=opts_per_line, char_limit=opt_char_limit)

    opts_html  = ""
    all_opts   = [o for g in option_groups for o in g]

    is_koot = any(
        re.search(r'[①②③④⑤⑥⑦⑧]', o['text']) or
        re.search(r'\b[ivxIVX]{1,4}\s*[-–]\s*[ivxIVX]{1,4}', o['text']) or
        re.search(r'\b[ivxIVX]{1,4}\s*[-–]\s*[a-dA-D]', o['text']) or
        re.search(r'\(\s*[ivxIVX]{1,4}\s*\)\s*[-–]', o['text'])
        for o in all_opts
    ) if all_opts else False

    total_rows = len(option_groups)
    n_cols     = max(len(g) for g in option_groups) if option_groups else 1

    if is_koot:
        for idx, opt in enumerate(all_opts):
            is_last = (idx == len(all_opts) - 1)
            ans_span = (
                f"<span style='font-weight:900;margin-left:8px;'>{q['correct']}</span>"
                if is_last and show_correct_inline and q['correct'] else ""
            )
            opts_html += (
                f"<div style='margin-left:{l2px + 10}px;font-size:{opt_font}pt;"
                f"line-height:{line_spacing}pt;'>"
                f"{opt['key']} {opt['text']}{ans_span}</div>"
            )
    else:
        for row_idx, group in enumerate(option_groups):
            is_last = (row_idx == total_rows - 1)
            cells_html = "".join(
                f"<div style='min-width:0;overflow-wrap:break-word;word-break:break-word;"
                f"font-weight:{'700' if opt_bold else '400'};'>"
                f"{o['key']} {o['text']}</div>"
                for o in group
            )
            for _ in range(n_cols - len(group)):
                cells_html += "<div></div>"

            if show_correct_inline and is_last:
                opts_html += (
                    f"<div style='display:flex;align-items:center;margin-left:{l2px}px;'>"
                    f"  <div style='display:grid;grid-template-columns:repeat({n_cols},1fr);"
                    f"flex:1;font-size:{opt_font}pt;column-gap:6px;"
                    f"line-height:{line_spacing}pt;'>{cells_html}</div>"
                    f"  <span style='font-weight:900;font-size:{opt_font+1.5}pt;"
                    f"white-space:nowrap;margin-left:8px;flex-shrink:0;'>{q.get('correct','')}</span>"
                    f"</div>"
                )
            else:
                opts_html += (
                    f"<div style='display:grid;grid-template-columns:repeat({n_cols},1fr);"
                    f"margin-left:{l2px}px;font-size:{opt_font}pt;column-gap:6px;"
                    f"line-height:{line_spacing}pt;'>{cells_html}</div>"
                )

    # ── Kathan lines HTML ──────────────────────────────────────────────────
    kathan_html = ""
    if q.get('kathan_lines'):
        bg_style = "background:#F5F5F5;border-radius:3px;padding:1px 4px;" if kathan_bg else ""
        for line in q['kathan_lines']:
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            kathan_html += (
                f"<div style='margin-left:{kl_px}px;{bg_style}"
                f"font-size:{kathan_font}pt;font-weight:{'bold' if kathan_bold else 'normal'};"
                f"line-height:{line_spacing}pt;'>{safe_line}</div>"
            )

    # ── Suchi / match table HTML ───────────────────────────────────────────
    suchi_html = ""
    if q.get('suchi_rows'):
        h1_lbl, h2_lbl = q.get('suchi_col_headers', ("सूची-I", "सूची-II"))
        suchi_html += (
            f"<table style='margin-left:{l2px}px;border-collapse:collapse;"
            f"font-size:{q_font}pt;width:calc(100% - {l2px}px);table-layout:fixed;'>"
            f"<colgroup><col style='width:50%;'><col style='width:50%;'></colgroup>"
            f"<tr><th style='background:#D9D9D9;padding:2px 6px;text-align:left;'>{h1_lbl}</th>"
            f"<th style='background:#D9D9D9;padding:2px 6px;text-align:left;'>{h2_lbl}</th></tr>"
        )
        for left, right in q['suchi_rows']:
            suchi_html += (
                f"<tr>"
                f"<td style='width:50%;padding:1px 6px 1px 0;vertical-align:top;word-wrap:break-word;'>{left}</td>"
                f"<td style='width:50%;padding:1px 0 1px 6px;vertical-align:top;word-wrap:break-word;'>{right}</td>"
                f"</tr>"
            )
        suchi_html += "</table>"

    if q.get('match_following_rows'):
        lp  = mtf_left_pct
        rp  = 100 - lp
        bdr = 'border-right:1px solid #ccc;' if mtf_inside_v else ''
        suchi_html += (
            f"<table style='margin-left:{l2px}px;border-collapse:collapse;"
            f"font-size:{q_font}pt;width:calc(100% - {l2px}px);table-layout:fixed;'>"
            f"<colgroup><col style='width:{lp}%;'><col style='width:{rp}%;'></colgroup>"
        )
        if mtf_show_header:
            suchi_html += (
                f"<tr>"
                f"<th style='background:#D9D9D9;padding:2px 6px;{bdr}'>{mtf_h1}</th>"
                f"<th style='background:#D9D9D9;padding:2px 6px;'>{mtf_h2}</th>"
                f"</tr>"
            )
        for left, right in q['match_following_rows']:
            suchi_html += (
                f"<tr>"
                f"<td style='padding:2px 6px;vertical-align:top;{bdr}'>{left}</td>"
                f"<td style='padding:2px 6px;vertical-align:top;'>{right}</td>"
                f"</tr>"
            )
        suchi_html += "</table>"

    kg = q.get('koot_grid', {})
    if kg.get('is_grid'):
        col_headers = kg.get('col_headers', [])
        grid_rows   = kg.get('rows', [])
        n_val       = max((len(r[1]) for r in grid_rows), default=0)
        correct_key = (q.get('correct') or "").strip("() ").upper()
        suchi_html += (
            f"<table style='margin-left:{l2px}px;border-collapse:collapse;"
            f"font-size:{opt_font}pt;margin-top:2px;'>"
        )
        if col_headers:
            header_cells = "".join(
                f"<th style='padding:1px 6px;text-align:center;"
                f"background:#D9D9D9;border:1px solid #ccc;'>{h}</th>"
                for h in col_headers[:n_val]
            )
            suchi_html += (
                f"<tr>"
                f"<th style='padding:1px 6px;background:#D9D9D9;border:1px solid #ccc;'></th>"
                f"{header_cells}</tr>"
            )
        for key_label, val_list in grid_rows:
            is_correct = key_label.strip("() ").upper() == correct_key
            fw = "font-weight:bold;" if is_correct else ""
            bg = "background:#FFF9C4;" if is_correct else ""
            val_cells = "".join(
                f"<td style='padding:1px 8px;text-align:center;{fw}{bg}"
                f"border:1px solid #eee;'>{val_list[i] if i < len(val_list) else ''}</td>"
                for i in range(n_val)
            )
            suchi_html += (
                f"<tr><td style='padding:1px 6px;{fw}{bg}border:1px solid #eee;'>"
                f"{key_label}</td>{val_cells}</tr>"
            )
        suchi_html += "</table>"
        opts_html = ""

    expl_html = ""
    if q['explanation'] or q.get('explanation_images'):
        _label   = expl_prefix if expl_prefix.strip() else "व्याख्या"
        _prefix  = f"➤ {_label}: " if expl_bullet else f"{_label}: "
        bg_style = "background-color:#F0F0F0;padding:2px 4px;border-radius:3px;" if expl_bg else ""
        expl_html += (
            f"<div style='margin-left:{l2px}px;{bg_style}font-size:{expl_font}pt;'>"
            f"<span style='font-weight:bold;'>{_prefix}</span>"
        )
        if q['explanation']:
            expl_html += q['explanation'].replace('|', '<br>')
        expl_html += "</div>"
        for img_bytes, _, __ in q.get('explanation_images', []):
            b64        = base64.b64encode(img_bytes).decode()
            expl_html += (
                f'<div style="margin-left:{l2px}px;">'
                f'<img src="data:image/png;base64,{b64}" style="max-width:100%;height:auto;"></div>'
            )

    question_html = q['question'].replace('\n', '<br>')
    q_html = (
        f"<div style='margin-left:{l2px}px;text-indent:{l1px - l2px}px;"
        f"font-size:{q_font}pt;font-weight:{'700' if q_bold else '400'};"
        f"margin-bottom:2px;white-space:pre-wrap;'>"
        f"{q['no']}. {question_html}</div>"
    )
    meta_html = ""
    if include_metadata and q.get('metadata'):
        meta_html = (
            f"<div style='text-align:right;font-size:{max(q_font-1,5)}pt;"
            f"font-style:italic;margin-left:{l2px}px;color:#555;'>"
            f"{q['metadata']}</div>"
        )

    return f"""
<div class="qblock">
  {q_html}
  {meta_html}
  {kathan_html}
  {suchi_html}
  {opts_html}
  {expl_html}
  {('<hr>' if show_separator else '')}
</div>"""


def build_preview_with_pagination(questions, q_per_page, heading_text):
    total_pages = (len(questions) + q_per_page - 1) // q_per_page
    pages_html  = []
    font_stack = ("'Arial','Calibri','Georgia',sans-serif"
                  if st.session_state.get("_is_english_doc")
                  else "'Mangal','Nirmala UI','Noto Sans Devanagari','Arial',sans-serif")
    page_label = "Page" if st.session_state.get("_is_english_doc") else "पृष्ठ"
    for page_num in range(1, total_pages + 1):
        start        = (page_num - 1) * q_per_page
        end          = min(start + q_per_page, len(questions))
        content_html = "".join(render_q_preview(q) for q in questions[start:end])
        pages_html.append(f"""
<div class="page" style="width:{page_width*96}px;min-height:{page_height*96}px;background:white;
  margin:0 auto 20px auto;padding:{top_margin*96}px {right_margin*96}px {bottom_margin*96}px {left_margin*96}px;
  box-shadow:0 4px 24px rgba(0,0,0,0.5);">
  <div style="background:#E6E6E6;padding:4px;border-radius:3px;text-align:center;
    font-weight:bold;margin-bottom:10px;">
    {chapter_heading if chapter_heading.strip() else heading_text} &nbsp;&nbsp;|&nbsp;&nbsp; {page_label} {page_num}
  </div>
  <div style="column-count:{num_columns};column-gap:18px;">{content_html}</div>
</div>""")

    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  *{{box-sizing:border-box;margin:0;padding:0;}}
  body{{background:#666;font-family:{font_stack};padding:20px;}}
  .qblock{{margin-bottom:5px;padding-bottom:4px;break-inside:avoid;page-break-inside:avoid;}}
  hr{{margin:4px 0;border:0;border-top:1px dotted #ccc;}}
</style>
</head><body>{''.join(pages_html)}</body></html>"""


# =============================================================================
# PDF GENERATION  (unchanged — kathan_lines rendered as paragraphs)
# =============================================================================
def register_devanagari_font():
    for path in [
        "C:/Windows/Fonts/Mangal.ttf",
        "C:/Windows/Fonts/Nirmala.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Mangal.ttf",
        "/usr/share/fonts/truetype/lohit/Lohit-Devanagari.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
    ]:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont('Devanagari', path))
                return 'Devanagari'
            except Exception:
                continue
    return 'Helvetica'


def generate_pdf(questions, chapter_title):
    font   = register_devanagari_font()
    buffer = BytesIO()
    doc_pdf = SimpleDocTemplate(
        buffer,
        pagesize=(page_width * inch, page_height * inch),
        topMargin=top_margin * inch, bottomMargin=bottom_margin * inch,
        leftMargin=left_margin * inch, rightMargin=right_margin * inch
    )
    styles = getSampleStyleSheet()
    l1 = level1_indent * inch
    l2 = level2_indent * inch
    kl = (level2_indent + kathan_indent_extra) * inch

    sQ     = ParagraphStyle('Q',  parent=styles['Normal'], fontSize=q_font,      leading=line_spacing,
                             fontName=font, spaceAfter=para_spacing, leftIndent=l2, firstLineIndent=l1-l2)
    sMeta  = ParagraphStyle('M',  parent=styles['Normal'], fontSize=6,            leading=line_spacing,
                             fontName=font, alignment=TA_RIGHT, spaceAfter=para_spacing, leftIndent=l2)
    sOpt   = ParagraphStyle('O',  parent=styles['Normal'], fontSize=opt_font,     leading=line_spacing,
                             fontName=font, spaceAfter=para_spacing, leftIndent=l2)
    sAns   = ParagraphStyle('A',  parent=styles['Normal'], fontSize=opt_font+1.5, leading=line_spacing,
                             fontName=font, alignment=TA_RIGHT, spaceAfter=para_spacing, leftIndent=l2)
    sExpl  = ParagraphStyle('E',  parent=styles['Normal'], fontSize=expl_font,    leading=line_spacing,
                             fontName=font, spaceAfter=para_spacing*2, leftIndent=l2, firstLineIndent=l1-l2,
                             backColor=colors.HexColor('#F0F0F0') if expl_bg else None)
    sH     = ParagraphStyle('H',  parent=styles['Normal'], fontSize=header_font,  leading=header_font+2,
                             fontName=font, alignment=TA_CENTER,
                             backColor=colors.HexColor('#E6E6E6') if header_bg else None, spaceAfter=6)
    sSuchi = ParagraphStyle('S',  parent=styles['Normal'], fontSize=q_font,       leading=line_spacing,
                             fontName=font, leftIndent=l2)
    sKathan = ParagraphStyle('K', parent=styles['Normal'], fontSize=kathan_font,  leading=line_spacing,
                              fontName=font, leftIndent=kl,
                              backColor=colors.HexColor('#F5F5F5') if kathan_bg else None)

    story = [Paragraph(f"{book_name} — {chapter_title}", sH)]

    sSep = ParagraphStyle('Sep', parent=styles['Normal'], fontSize=header_font, leading=header_font+2,
                           fontName=font, alignment=TA_CENTER,
                           backColor=colors.HexColor('#3B3B3B'), textColor=colors.white, spaceAfter=6, spaceBefore=6)

    for q in questions:
        if q.get('is_separator'):
            safe_name = (q.get('text', '') or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(f"<b>📄 {safe_name}</b>", sSep))
            continue

        story.append(Paragraph(f"<b>{q['no']}.</b> {q['question']}", sQ))

        if include_metadata and q.get('metadata'):
            story.append(Paragraph(q['metadata'], sMeta))

        # Kathan lines
        for line in q.get('kathan_lines', []):
            safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(
                f"<b>{safe}</b>" if kathan_bold else safe, sKathan
            ))

        if q.get('suchi_rows'):
            content_w = page_width - left_margin - right_margin
            col_gap   = 0.08 if num_columns == 3 else 0.12
            col_w     = (content_w - col_gap * (num_columns - 1)) / num_columns
            half_w    = (col_w - level2_indent) * inch / 2
            h1_lbl, h2_lbl = q.get('suchi_col_headers', ("सूची-I", "सूची-II"))
            tdata = [[Paragraph(f"<b>{h1_lbl}</b>", sSuchi), Paragraph(f"<b>{h2_lbl}</b>", sSuchi)]]
            tdata += [[Paragraph(left, sSuchi), Paragraph(right, sSuchi)]
                      for left, right in q['suchi_rows']]
            t = Table(tdata, colWidths=[half_w, half_w])
            t.setStyle(TableStyle([
                ('BACKGROUND',    (0,0), (-1,0),  colors.HexColor('#D9D9D9')),
                ('VALIGN',        (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING',   (0,0), (-1,-1), 2),
                ('RIGHTPADDING',  (0,0), (-1,-1), 2),
                ('TOPPADDING',    (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('GRID',          (0,0), (-1,-1), 0, colors.white),
            ]))
            story.append(t)

        if q.get('match_following_rows'):
            content_w = page_width - left_margin - right_margin
            col_gap   = 0.08 if num_columns == 3 else 0.12
            col_w     = (content_w - col_gap * (num_columns - 1)) / num_columns
            avail_w   = (col_w - level2_indent) * inch
            lw        = avail_w * (mtf_left_pct / 100.0)
            rw        = avail_w - lw
            mdata     = [[Paragraph(left, sSuchi), Paragraph(right, sSuchi)]
                         for left, right in q['match_following_rows']]
            if mtf_show_header:
                mdata.insert(0, [Paragraph(f"<b>{mtf_h1}</b>", sSuchi),
                                 Paragraph(f"<b>{mtf_h2}</b>", sSuchi)])
            mt = Table(mdata, colWidths=[lw, rw])
            ts_cmds = [
                ('VALIGN',        (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING',   (0,0), (-1,-1), 2),
                ('RIGHTPADDING',  (0,0), (-1,-1), 2),
                ('TOPPADDING',    (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('GRID',          (0,0), (-1,-1), 0, colors.white),
            ]
            if mtf_show_header:
                ts_cmds.append(('BACKGROUND', (0,0), (-1,0), colors.HexColor('#D9D9D9')))
            if mtf_inside_v:
                ts_cmds.append(('LINEAFTER', (0,0), (0,-1), 0.5, colors.HexColor('#AAAAAA')))
            mt.setStyle(TableStyle(ts_cmds))
            story.append(mt)

        kg = q.get('koot_grid', {})
        if not kg.get('is_grid'):
            opt_groups = q.get('_layout', layout_options(
                q['options'], max_per_line=opts_per_line, char_limit=opt_char_limit
            ))
            if opt_groups:
                content_w  = page_width - left_margin - right_margin
                col_gap    = 0.08 if num_columns == 3 else 0.12
                col_w      = (content_w - col_gap * (num_columns - 1)) / num_columns
                ans_col_w  = 0.38 * inch
                opt_avail  = (col_w - level2_indent) * inch - ans_col_w

                for row_idx, group in enumerate(opt_groups):
                    is_last    = row_idx == len(opt_groups) - 1
                    n_cols_row = len(group)
                    opt_col_w  = opt_avail / n_cols_row
                    row_data   = [Paragraph(f"{o['key']} {o['text']}", sOpt) for o in group]
                    if is_last and show_correct_inline and q['correct']:
                        row_data.append(Paragraph(f"<b>{q['correct']}</b>", sAns))
                    else:
                        row_data.append(Paragraph("", sOpt))
                    col_widths_pdf = [opt_col_w] * n_cols_row + [ans_col_w]
                    t = Table([row_data], colWidths=col_widths_pdf)
                    t.setStyle(TableStyle([
                        ('VALIGN',        (0,0), (-1,-1), 'TOP'),
                        ('LEFTPADDING',   (0,0), (-1,-1), 2),
                        ('RIGHTPADDING',  (0,0), (-1,-1), 2),
                        ('TOPPADDING',    (0,0), (-1,-1), 0),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                        ('GRID',          (0,0), (-1,-1), 0, colors.white),
                    ]))
                    story.append(t)

        if q['explanation'] or q.get('explanation_images'):
            _label    = expl_prefix if expl_prefix.strip() else "व्याख्या"
            _heading  = (f"• {_label} : " if expl_bullet else f"{_label} : ")
            expl_text = _heading + (q['explanation'] if q['explanation'] else "")
            story.append(Paragraph(expl_text.replace('|', '<br/>'), sExpl))

        if show_separator:
            story.append(Spacer(1, 2))

    doc_pdf.build(story)
    buffer.seek(0)
    return buffer


# =============================================================================
# CHAPTER TITLE EXTRACTION  (unchanged)
# =============================================================================
def extract_chapter_title_from_lines(lines):
    for t in lines[:10]:
        t = (t or "").strip()
        if not t:
            continue
        if "अध्याय" in t or "CHAPTER" in t.upper():
            return t[:80] + ("..." if len(t) > 80 else "")
        if (t.isupper() or t.istitle()) and len(t) < 80 and not re.match(r'^\d', t):
            return t
    return "RBD PUBLICATION — Chapter"


def extract_chapter_title(doc):
    return extract_chapter_title_from_lines([p.text for p in doc.paragraphs[:10]])


# =============================================================================
# NUMBERING MODES  (Default / Sequenced / Customized)
# =============================================================================
def apply_numbering_mode(questions, mode, start_number=1):
    """
    mode: 'default'    -> keep the number exactly as it appears in the source file
          'sequenced'  -> renumber 1..N continuously across ALL merged papers
          'customized' -> renumber start_number..start_number+N-1 continuously
    Separator markers (is_separator=True) are skipped — they don't consume a number.
    """
    counter = start_number
    for q in questions:
        if q.get('is_separator'):
            continue
        if mode == 'default':
            q['no'] = q.get('orig_no', q.get('no', ''))
        else:  # 'sequenced' or 'customized'
            q['no'] = str(counter)
            counter += 1
    return questions


# =============================================================================
# MULTI-FILE MERGE — filename separator marker
# =============================================================================
def _make_separator_question(filename):
    return {
        "no": "", "orig_no": "", "question": "",
        "kathan_lines": [], "suchi_rows": [],
        "suchi_col_headers": ("सूची-I", "सूची-II"),
        "match_following_rows": [], "koot_grid": {"is_grid": False},
        "options": [], "correct": "", "explanation": "",
        "explanation_images": [], "metadata": "", "_layout": None,
        "is_separator": True, "text": filename,
    }


# =============================================================================
# FILENAME SLUGIFY  (for downloads)
# =============================================================================
def slugify_filename(text, default="Formatted_Output"):
    text = (text or "").strip()
    if not text:
        text = default
    text = re.sub(r'\*+', '', text)
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'[\\/:*?"<>|]', '', text)     # illegal filename chars
    text = re.sub(r'\s+', '_', text)
    text = text.strip('_.')
    return text if text else default


# =============================================================================
# JSON / TXT INPUT PARSERS  (for the multi-format uploader)
# =============================================================================
def _coerce_json_options(raw_options):
    options = []
    if not raw_options:
        return options
    default_keys = ['(a)', '(b)', '(c)', '(d)', '(e)', '(f)']
    for i, opt in enumerate(raw_options):
        if isinstance(opt, dict):
            key  = opt.get('key')  or (default_keys[i] if i < len(default_keys) else f"({i+1})")
            text = opt.get('text', '')
        else:
            key  = default_keys[i] if i < len(default_keys) else f"({i+1})"
            text = str(opt)
        options.append({"key": key, "text": text})
    return options


def _json_item_to_question(item, fallback_no):
    no = str(item.get('no', fallback_no))
    suchi_rows = [tuple(r) for r in item.get('suchi_rows', [])]
    mtf_rows   = [tuple(r) for r in item.get('match_following_rows', [])]
    headers    = item.get('suchi_col_headers', ["सूची-I", "सूची-II"])
    return {
        "no": no, "orig_no": no,
        "question": item.get('question', ''),
        "kathan_lines": item.get('kathan_lines', []),
        "suchi_rows": suchi_rows,
        "suchi_col_headers": tuple(headers) if len(headers) == 2 else ("सूची-I", "सूची-II"),
        "match_following_rows": mtf_rows,
        "koot_grid": item.get('koot_grid', {"is_grid": False}),
        "options": _coerce_json_options(item.get('options')),
        "correct": item.get('correct') or item.get('answer', ''),
        "explanation": item.get('explanation', ''),
        "explanation_images": [],
        "metadata": item.get('metadata', ''),
        "_layout": None,
    }


def _read_upload_text(file_obj):
    raw = file_obj.read()
    if isinstance(raw, str):
        return raw
    try:
        return raw.decode('utf-8')
    except UnicodeDecodeError:
        return raw.decode('latin-1', errors='ignore')


def parse_txt_file(file_obj):
    """TXT files must use the same <question>/<option>/<answer>/<expl>/<kathan>/<suchi>
    tag format as the tagged DOCX files."""
    text = _read_upload_text(file_obj)

    if not _has_tags(text):
        st.warning(
            f"⚠️ **{file_obj.name}** — no `<question>`/`<option>`/... tags found. "
            "TXT files must use the same tag format as your tagged DOCX files. Skipping this file."
        )
        return [], False, file_obj.name

    questions  = parse_questions_tagged(text, is_text=True)
    is_english = not any('\u0900' <= c <= '\u097F' for c in text)
    title      = extract_chapter_title_from_lines(text.split('\n'))
    return questions, is_english, title


def parse_json_file(file_obj):
    """JSON files: either a plain list of question objects, or
    {"chapter_title": "...", "questions": [...]}.
    Each question object supports: no, question, options, correct/answer,
    explanation, kathan_lines, suchi_rows, match_following_rows, koot_grid, metadata."""
    text = _read_upload_text(file_obj)
    try:
        data = json.loads(text)
    except Exception as e:
        st.error(f"❌ **{file_obj.name}** — invalid JSON: {e}. Skipping this file.")
        return [], False, file_obj.name

    if isinstance(data, dict):
        title = data.get('chapter_title') or data.get('title') or file_obj.name
        items = data.get('questions', [])
    else:
        title = file_obj.name
        items = data

    questions = [_json_item_to_question(item, idx + 1) for idx, item in enumerate(items)]
    combined  = " ".join(q['question'] for q in questions)
    is_english = not any('\u0900' <= c <= '\u097F' for c in combined)
    return questions, is_english, title


# =============================================================================
# .DOC OUTPUT VIA LIBREOFFICE  (best-effort — needs `soffice` on PATH)
# =============================================================================
def convert_docx_bytes_to_doc(docx_bytes):
    """Converts DOCX bytes to legacy .doc bytes using headless LibreOffice.
    Returns (doc_bytes_or_None, error_message_or_None)."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None, (
            "LibreOffice (`soffice`) isn't installed on this host, so a true .doc file "
            "can't be produced. Add a `packages.txt` file containing `libreoffice` to your "
            "repo root if you're on Streamlit Community Cloud, or install LibreOffice on "
            "your server, then redeploy."
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            result = subprocess.run(
                [soffice, "--headless", "--norestore", "--convert-to", "doc", "--outdir", tmpdir, docx_path],
                capture_output=True, timeout=120
            )
        except Exception as e:
            return None, f"LibreOffice conversion failed to run: {e}"

        doc_path = os.path.join(tmpdir, "input.doc")
        if not os.path.exists(doc_path):
            err = result.stderr.decode(errors='ignore') if result.stderr else "unknown error"
            return None, f"LibreOffice conversion did not produce an output file: {err}"

        with open(doc_path, "rb") as f:
            return f.read(), None


def parse_uploaded_file(file_obj):
    """Universal per-file parser dispatcher: routes to the right parser
    based on file extension. Returns (questions, is_english, chapter_title_guess)."""
    ext = file_obj.name.rsplit('.', 1)[-1].lower() if '.' in file_obj.name else ''
    if ext == 'docx':
        doc = Document(file_obj)
        questions, is_english = parse_questions(doc)
        title = extract_chapter_title(doc)
        return questions, is_english, title
    elif ext == 'txt':
        return parse_txt_file(file_obj)
    elif ext == 'json':
        return parse_json_file(file_obj)
    else:
        st.warning(f"⚠️ **{file_obj.name}** — unsupported file type `.{ext}`. Skipping.")
        return [], False, file_obj.name


# =============================================================================
# MAIN APP
# =============================================================================
if uploaded_files_ordered and upload_mode == "Batch (process each file separately)":
    # -------------------------------------------------------------------
    # BATCH MODE — each file parsed, numbered, and generated independently.
    # Same global settings (page design / fonts / numbering mode) apply to
    # every file; only the output name differs per file.
    # -------------------------------------------------------------------
    with st.spinner(f"Parsing {len(uploaded_files_ordered)} file(s) independently..."):
        batch_results = []
        used_names = set()
        for f in uploaded_files_ordered:
            qs, is_eng, title = parse_uploaded_file(f)
            qs = apply_numbering_mode(qs, numbering_mode, numbering_start)
            k = _file_key(f)
            custom_name  = (batch_output_names.get(k) or f.name.rsplit('.', 1)[0]).strip()
            base_out     = slugify_filename(custom_name, default=f"Output_{len(batch_results) + 1}")
            out_name, suf = base_out, 2
            while out_name in used_names:
                out_name = f"{base_out}_{suf}"
                suf += 1
            used_names.add(out_name)
            batch_results.append({
                "source_name": f.name, "questions": qs, "is_english": is_eng,
                "title": title, "out_name": out_name,
            })

    total_qs = sum(len(r["questions"]) for r in batch_results)
    st.success(f"✅ {len(batch_results)} file(s) parsed independently — {total_qs} questions total.")

    docx_format_choice = st.radio(
        "DOCX output format (applies to all files)",
        ["Modern (.docx)", "Legacy (.doc) — via LibreOffice"],
        horizontal=True, key="batch_docx_format_choice"
    )

    def _build_docx_bytes_for(r):
        precompute_layouts(r["questions"])
        d   = generate_multi_page_docx(r["questions"], r["title"])
        buf = BytesIO()
        d.save(buf)
        buf.seek(0)
        return buf.getvalue()

    file_tabs = st.tabs([f"{i + 1}. {r['out_name']}" for i, r in enumerate(batch_results)])
    for tab, r in zip(file_tabs, batch_results):
        with tab:
            lang_tag = "🇬🇧 English" if r["is_english"] else "🇮🇳 Hindi"
            st.write(f"Source: **{r['source_name']}**  |  {len(r['questions'])} question(s)  |  {lang_tag}")
            colA, colB = st.columns(2)
            with colA:
                if st.button("🚀 Generate DOCX", key=f"batch_gen_docx_{r['out_name']}"):
                    with st.spinner(f"Generating {r['out_name']}..."):
                        docx_bytes = _build_docx_bytes_for(r)
                        if docx_format_choice.startswith("Legacy"):
                            doc_bytes, err = convert_docx_bytes_to_doc(docx_bytes)
                            if doc_bytes:
                                st.download_button(
                                    "📥 Download DOC", doc_bytes, f"{r['out_name']}.doc",
                                    "application/msword", key=f"batch_dl_doc_{r['out_name']}"
                                )
                            else:
                                st.error(f"❌ Couldn't produce a real .doc file: {err}")
                                st.download_button(
                                    "📥 Download DOCX instead", docx_bytes, f"{r['out_name']}.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"batch_dl_docx_{r['out_name']}"
                                )
                        else:
                            st.download_button(
                                "📥 Download DOCX", docx_bytes, f"{r['out_name']}.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"batch_dl_docx_{r['out_name']}"
                            )
                        st.success("🎉 Ready!")
            with colB:
                if st.button("📑 Generate PDF", key=f"batch_gen_pdf_{r['out_name']}"):
                    with st.spinner(f"Generating {r['out_name']}.pdf..."):
                        pdf_buffer = generate_pdf(r["questions"], r["title"])
                        st.download_button(
                            "📥 Download PDF", pdf_buffer, f"{r['out_name']}.pdf",
                            "application/pdf", key=f"batch_dl_pdf_{r['out_name']}"
                        )
                        st.success("🎉 Ready!")

    st.markdown("---")
    if st.button("📦 Generate All & Download ZIP", key="batch_zip_btn"):
        with st.spinner("Generating all files and zipping..."):
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for r in batch_results:
                    docx_bytes = _build_docx_bytes_for(r)
                    if docx_format_choice.startswith("Legacy"):
                        doc_bytes, err = convert_docx_bytes_to_doc(docx_bytes)
                        if doc_bytes:
                            zf.writestr(f"{r['out_name']}.doc", doc_bytes)
                        else:
                            zf.writestr(f"{r['out_name']}.docx", docx_bytes)  # fallback if LibreOffice missing
                    else:
                        zf.writestr(f"{r['out_name']}.docx", docx_bytes)
            zip_buffer.seek(0)
            st.download_button(
                "📥 Download All as ZIP", zip_buffer, "Batch_Output.zip",
                "application/zip", key="batch_zip_dl"
            )
            st.success("🎉 ZIP ready!")

elif uploaded_files_ordered:
    with st.spinner(f"Parsing {len(uploaded_files_ordered)} file(s)..."):
        per_file_results = []   # list of (filename, questions, is_english, title)
        for f in uploaded_files_ordered:
            qs, is_eng, title = parse_uploaded_file(f)
            per_file_results.append((f.name, qs, is_eng, title))

        questions = []
        for i, (fname, qs, is_eng, title) in enumerate(per_file_results):
            if i > 0 and add_filename_separator:
                questions.append(_make_separator_question(fname))
            questions.extend(qs)

        questions = apply_numbering_mode(questions, numbering_mode, numbering_start)

        _is_english_doc = per_file_results[0][2] if per_file_results else False
        chapter_title    = per_file_results[0][3] if per_file_results else "RBD PUBLICATION — Chapter"
        st.session_state["_is_english_doc"] = _is_english_doc

    lang_tag    = "🇬🇧 English" if _is_english_doc else "🇮🇳 Hindi"
    real_qs     = [q for q in questions if not q.get('is_separator')]
    has_kathan  = sum(1 for q in real_qs if q.get('kathan_lines'))
    has_suchi   = sum(1 for q in real_qs if q.get('suchi_rows'))
    st.success(
        f"✅ {len(real_qs)} questions parsed from {len(uploaded_files_ordered)} file(s)!  ({lang_tag} paper)  "
        f"| Suchi: {has_suchi}  | Kathan: {has_kathan}"
    )
    if len(per_file_results) > 1:
        with st.expander("📊 Per-file breakdown"):
            for fname, qs, is_eng, title in per_file_results:
                st.write(f"**{fname}** — {len(qs)} question(s)")

    _init_table_editor_state(questions)
    if not chapter_heading.strip():
        chapter_heading = re.sub(r'\*+', '', chapter_title).strip()

    if auto_fill:
        sample_size      = min(10, len(questions))
        total_lines      = sum(estimate_q_lines(q) for q in questions[:sample_size])
        avg_lines        = total_lines / sample_size if sample_size > 0 else 10
        usable_height    = page_height - top_margin - bottom_margin - 1.2
        lines_per_page   = usable_height / (line_spacing / 72.0)
        q_per_page_est   = max(1, int(lines_per_page / avg_lines))
        total_pages_est  = (len(questions) + q_per_page_est - 1) // q_per_page_est
    else:
        q_per_page_est  = 20
        total_pages_est = (len(questions) + q_per_page_est - 1) // q_per_page_est

    st.info(f"📄 Estimated pages: {total_pages_est} ({'auto' if auto_fill else 'fixed'})")

    tab1, tab2, tab3 = st.tabs(["📄 Page Preview", "🔍 Parsed Data", "🔧 Manual Table Editor"])
    with tab1:
        preview_html = build_preview_with_pagination(
            questions, q_per_page_est,
            chapter_heading if chapter_heading.strip() else chapter_title
        )
        st.components.v1.html(preview_html, height=1200, scrolling=True)
    with tab2:
        shown = 0
        for q in questions:
            if q.get('is_separator'):
                st.markdown(f"**📄 — New paper begins: {q.get('text', '')} —**")
                continue
            if shown >= 5:
                break
            shown += 1
            with st.expander(f"Q{q['no']} – {q['question'][:60]}…"):
                st.write("**Options:**",             q.get('options', []))
                st.write("**Correct Answer:**",      q.get('correct', ''))
                st.write("**Explanation:**",         (q.get('explanation') or '')[:500])
                st.write(f"**Kathan lines:**         {len(q.get('kathan_lines', []))}",
                         q.get('kathan_lines', []))
                st.write(f"**Suchi rows:**           {len(q.get('suchi_rows', []))}",
                         q.get('suchi_rows', []))
                st.write(f"**Match-following rows:** {len(q.get('match_following_rows', []))}",
                         q.get('match_following_rows', []))
    with tab3:
        render_table_editor_tab(st.session_state["te_questions"])

    st.markdown("---")
    output_filename_base = slugify_filename(
        chapter_heading if chapter_heading.strip() else chapter_title
    )

    c1, c2 = st.columns(2)
    with c1:
        docx_format_choice = st.radio(
            "DOCX output format", ["Modern (.docx)", "Legacy (.doc) — via LibreOffice"],
            horizontal=True, key="docx_format_choice"
        )
        if st.button("🚀 Generate DOCX"):
            with st.spinner(f"Generating DOCX with font: {FONT_DOCX}..."):
                precompute_layouts(st.session_state["te_questions"])
                final_doc   = generate_multi_page_docx(st.session_state["te_questions"], chapter_title)
                docx_buffer = BytesIO()
                final_doc.save(docx_buffer)
                docx_buffer.seek(0)

                if docx_format_choice.startswith("Legacy"):
                    doc_bytes, err = convert_docx_bytes_to_doc(docx_buffer.getvalue())
                    if doc_bytes:
                        st.download_button(
                            "📥 Download DOC",
                            doc_bytes,
                            f"{output_filename_base}.doc",
                            "application/msword"
                        )
                        st.success(f"🎉 DOC ready! Font used: **{FONT_DOCX}**")
                    else:
                        st.error(f"❌ Couldn't produce a real .doc file: {err}")
                        st.download_button(
                            "📥 Download DOCX instead",
                            docx_buffer,
                            f"{output_filename_base}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.download_button(
                        "📥 Download DOCX",
                        docx_buffer,
                        f"{output_filename_base}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success(f"🎉 DOCX ready! Font used: **{FONT_DOCX}**")
    with c2:
        if st.button("📑 Preview PDF"):
            with st.spinner("Generating PDF preview..."):
                pdf_buffer = generate_pdf(st.session_state["te_questions"], chapter_title)
                pdf_b64    = base64.b64encode(pdf_buffer.getvalue()).decode()
                st.markdown(
                    f'<iframe src="data:application/pdf;base64,{pdf_b64}" '
                    f'width="100%" height="800" type="application/pdf"></iframe>',
                    unsafe_allow_html=True
                )
                st.download_button(
                    "📥 Download PDF", pdf_buffer,
                    file_name=f"{output_filename_base}.pdf", mime="application/pdf"
                )
                st.success("🎉 PDF preview ready!")